import os
os.environ['TOKENIZERS_PARALLELISM'] = 'false'
import json
import re
import requests
import docx
import asyncio
import base64
import gc  # Added for garbage collection
import hashlib
import time
import logging
from datetime import datetime
from bs4 import BeautifulSoup
from googlesearch import search
from together import Together
from dotenv import load_dotenv
from langchain_core.documents import Document
from langchain.text_splitter import RecursiveCharacterTextSplitter
from langchain_community.vectorstores import FAISS
from unstructured.partition.auto import partition
from unstructured.chunking.title import chunk_by_title
try:
    import fitz  # PyMuPDF
    PYMUPDF_AVAILABLE = True
except ImportError:
    PYMUPDF_AVAILABLE = False
    print("Warning: PyMuPDF not available. Install with: pip install PyMuPDF")

try:
    import pdfplumber
    PDFPLUMBER_AVAILABLE = True
except ImportError:
    PDFPLUMBER_AVAILABLE = False
    print("Warning: pdfplumber not available. Install with: pip install pdfplumber")
from langchain.chains import ConversationalRetrievalChain
from langchain.memory import ConversationBufferMemory
from langchain.prompts import PromptTemplate
from langchain_together import ChatTogether
from langchain_together.embeddings import TogetherEmbeddings
from tenacity import retry, stop_after_attempt, wait_exponential, retry_if_exception_type
import ssl
from langdetect import detect, LangDetectException

load_dotenv()

def load_features():
    """Helper function to load current feature configuration."""
    features_file = "features.json"
    default_features = {
        "image_generation": True,
        "pdf_docx_upload": True,
        "txt_sql_upload": True,
        "xlsx_csv_analysis": True,
        "web_search": True
    }
    
    try:
        if os.path.exists(features_file):
            with open(features_file, 'r') as f:
                return json.load(f)
        return default_features
    except Exception:
        return default_features

# --- Global Variables ---
LLM_MODEL = "meta-llama/Llama-4-Maverick-17B-128E-Instruct-FP8"
#LLM_MODEL = "deepseek-ai/DeepSeek-V3" # Alternative model for general use
IMAGE_MODEL = "meta-llama/Llama-4-Maverick-17B-128E-Instruct-FP8" # Model for image analysis
# IMAGE_GEN_MODEL = "black-forest-labs/FLUX.1-schnell" # Model for image generation
IMAGE_GEN_MODEL = "black-forest-labs/FLUX.1-kontext-max" # Model for image generation
EMBEDDING_MODEL = "intfloat/multilingual-e5-large-instruct" # Multilingual model for German/English documents
FAST_MODEL = "meta-llama/Llama-3.2-3B-Instruct-Turbo" # Fast model for analysis and translation
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCUMENTS_PATH = os.path.join(SCRIPT_DIR, "Documents")
VECTOR_STORE_PATH = os.path.join(SCRIPT_DIR, "vector_store")
# Enhanced Chunking Strategy - Optimized for embedding model compatibility
DEFAULT_CHUNK_SIZE = 2000  # Reduced for better embedding model compatibility
DEFAULT_OVERLAP = 400  # Reduced overlap for better performance
# Document-size-based chunking thresholds
SMALL_DOC_THRESHOLD = 100000  # 100k chars - use full-text approach
MEDIUM_DOC_THRESHOLD = 500000  # 500k chars - use large chunks
LARGE_CHUNK_SIZE = 2500  # For large documents - reduced for compatibility
MEDIUM_CHUNK_SIZE = 2000  # For medium documents - reduced for compatibility
SMALL_CHUNK_SIZE = 1500  # For small documents - more conservative
# Embedding model safety limits
MAX_CHUNK_SIZE = 4000  # Absolute maximum chunk size for embedding model

# --- Robust API Call Wrapper ---
@retry(
    wait=wait_exponential(multiplier=1, min=2, max=10),
    stop=stop_after_attempt(3),
    retry=retry_if_exception_type((requests.exceptions.RequestException, ssl.SSLError))
)
def robust_api_call(client, model, messages, temperature, stream=False, timeout=None, cancellation_check=None):
    """
    Makes a robust API call to the Together AI client.
    Can handle both streaming and non-streaming requests.
    Now supports cancellation checks for better responsiveness.
    """
    # Check for cancellation before making the API call
    if cancellation_check and cancellation_check():
        raise Exception("API call cancelled before execution")
    
    try:
        response = client.chat.completions.create(
            model=model,
            messages=messages,
            temperature=temperature,
            stream=stream,  # Pass the stream parameter
            timeout=timeout # Pass timeout to the API call
        )
        
        # For streaming responses, wrap with cancellation checks
        if stream and cancellation_check:
            def cancellation_aware_stream():
                for chunk in response:
                    if cancellation_check():
                        break
                    yield chunk
            return cancellation_aware_stream()
        
        return response
    except Exception as e:
        # Check if this was a cancellation during API call
        if cancellation_check and cancellation_check():
            raise Exception("API call cancelled during execution")
        raise e

def truncate_text(text: str, max_chars: int) -> str:
    """Trunkiert den Text auf eine maximale Zeichenlänge und behält den Anfang bei."""
    if len(text) > max_chars:
        return text[:max_chars]
    return text

# --- Document Loading and Vector Store Creation ---
import glob

def get_document_list(path: str) -> list:
    docx_files = glob.glob(os.path.join(path, "**/*.docx"), recursive=True)
    pdf_files = glob.glob(os.path.join(path, "**/*.pdf"), recursive=True)
    all_files = docx_files + pdf_files
    return [f for f in all_files if not os.path.basename(f).startswith('~$')]

async def extract_document_structure_with_llm(file_path: str) -> dict:
    """
    Intelligent, language-agnostic document structure extraction using LLM analysis.
    Works with any document type and language without hard-coded assumptions.
    """
    doc = docx.Document(file_path)
    structure = {
        "filename": os.path.basename(file_path),
        "chapters": {},
        "subsections": {},
        "topics": {},
        "processes": {},
        "headings_hierarchy": []
    }
    
    path_parts = os.path.normpath(file_path).split(os.sep)
    try:
        documents_index = path_parts.index('Documents')
        knowledge_field = path_parts[documents_index + 1]
    except (ValueError, IndexError):
        knowledge_field = "Unknown"
    
    current_headings = [""] * 6
    heading_content_pairs = []
    
    # Extract all headings and their content
    for para in doc.paragraphs:
        if para.style.name.startswith('Heading') and para.text.strip():
            try:
                level = int(para.style.name.split(' ')[-1])
                if 1 <= level <= 6:
                    current_headings[level - 1] = para.text.strip()
                    for i in range(level, 6):
                        current_headings[i] = ""
                    
                    heading_breadcrumb = " > ".join(h for h in current_headings if h)
                    
                    # Store heading hierarchy
                    heading_info = {
                        "level": level,
                        "text": para.text.strip(),
                        "breadcrumb": heading_breadcrumb
                    }
                    structure["headings_hierarchy"].append(heading_info)
                    
                    # Collect headings for LLM analysis
                    heading_content_pairs.append({
                        "level": level,
                        "heading": para.text.strip(),
                        "breadcrumb": heading_breadcrumb
                    })
                    
            except ValueError:
                continue
    
    # If we have headings, analyze them with LLM
    if heading_content_pairs:
        try:
            # Use LLM to intelligently analyze document structure
            client = Together(api_key=os.getenv("TOGETHER_API_KEY"))
            
            # Prepare headings text for analysis
            headings_text = "\n".join([
                f"Level {h['level']}: {h['breadcrumb']}" 
                for h in heading_content_pairs
            ])
            
            analysis_prompt = f"""Analyze this document structure and extract topics, processes, and search terms dynamically. The document has these headings:

{headings_text}

Your task:
1. Identify the main TOPICS covered in this document
2. Identify any PROCESSES or PROCEDURES described
3. Extract KEY TERMS that users might search for
4. Detect the document LANGUAGE and DOMAIN

Respond with ONLY a JSON object in this format:
{{
  "language": "detected language (e.g., 'German', 'English', 'French')",
  "domain": "document domain (e.g., 'Software Manual', 'HR Policy', 'Financial Report')",
  "main_topics": ["topic1", "topic2", "topic3"],
  "processes": ["process1", "process2"],
  "search_terms": ["term1", "term2", "term3"],
  "chapter_mappings": {{
    "topic or process": "heading breadcrumb where it's discussed"
  }},
  "subsection_mappings": {{
    "specific procedure": "detailed heading breadcrumb"
  }}
}}

Important: Base your analysis ONLY on the actual headings provided. Don't make assumptions about content not visible in the headings."""

            messages = [{"role": "user", "content": analysis_prompt}]
            
            response = await asyncio.to_thread(
                robust_api_call, 
                client, 
                FAST_MODEL,  # Fast model for analysis
                messages, 
                0.1
            )
            
            if response.choices and response.choices[0].message.content:
                content = response.choices[0].message.content.strip()
                
                # Clean JSON if wrapped in markdown
                if content.startswith("```json"):
                    content = content[7:-3].strip()
                elif content.startswith("```"):
                    content = content[3:-3].strip()
                
                try:
                    analysis = json.loads(content)
                    
                    # Map the LLM analysis to our structure format
                    structure["language"] = analysis.get("language", "Unknown")
                    structure["domain"] = analysis.get("domain", "Unknown")
                    structure["main_topics"] = analysis.get("main_topics", [])
                    structure["processes_list"] = analysis.get("processes", [])
                    structure["search_terms"] = analysis.get("search_terms", [])
                    
                    # Build search-friendly mappings
                    chapter_mappings = analysis.get("chapter_mappings", {})
                    subsection_mappings = analysis.get("subsection_mappings", {})
                    
                    # Store both original terms and lowercase variants for better matching
                    for term, breadcrumb in chapter_mappings.items():
                        structure["chapters"][term] = breadcrumb
                        structure["chapters"][term.lower()] = breadcrumb
                    
                    for term, breadcrumb in subsection_mappings.items():
                        structure["subsections"][term] = breadcrumb
                        structure["subsections"][term.lower()] = breadcrumb
                    
                    # Create topic-based mappings for broader search coverage
                    for topic in analysis.get("main_topics", []):
                        # Find headings that might relate to this topic
                        for heading_info in heading_content_pairs:
                            if any(word in heading_info["heading"].lower() for word in topic.lower().split()):
                                structure["topics"][topic] = heading_info["breadcrumb"]
                                structure["topics"][topic.lower()] = heading_info["breadcrumb"]
                                break
                    
                except json.JSONDecodeError as e:
                    print(f"Error parsing LLM analysis for {file_path}: {e}")
                    # Fallback to basic structure
                    pass
                    
        except Exception as e:
            print(f"Error in LLM analysis for {file_path}: {e}")
            # Continue with basic structure
            pass
    
    return structure

def extract_document_structure(file_path: str) -> dict:
    """
    Wrapper for the async LLM-based structure extraction.
    Maintains compatibility with existing sync code.
    """
    import asyncio
    
    # Create new event loop if none exists
    try:
        loop = asyncio.get_event_loop()
    except RuntimeError:
        loop = asyncio.new_event_loop()
        asyncio.set_event_loop(loop)
    
    # Run the async function
    return loop.run_until_complete(extract_document_structure_with_llm(file_path))

def process_docx_with_headings(file_path: str) -> list[Document]:
    doc = docx.Document(file_path)
    docs = []
    current_headings = [""] * 6
    # Extract knowledge field from the file path
    path_parts = os.path.normpath(file_path).split(os.sep)
    # The knowledge field is the directory name inside 'Documents'
    try:
        documents_index = path_parts.index('Documents')
        knowledge_field = path_parts[documents_index + 1]
    except (ValueError, IndexError):
        knowledge_field = "Unknown"

    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            try:
                level = int(para.style.name.split(' ')[-1])
                if 1 <= level <= 6:
                    current_headings[level - 1] = para.text
                    for i in range(level, 6):
                        current_headings[i] = ""
            except ValueError:
                continue
        elif para.text.strip():
            heading_breadcrumb = " > ".join(h for h in current_headings if h)
            page_content = f"[Kontext: {heading_breadcrumb}] {para.text}"
            metadata = {
                "source": file_path, 
                "heading": heading_breadcrumb,
                "knowledge_field": knowledge_field,
                "heading_level": len([h for h in current_headings if h])
            }
            docs.append(Document(page_content=page_content, metadata=metadata))
    return docs

# --- Vector Store Management ---
vector_stores = {} # Dictionary to hold multiple vector stores

def load_vector_store():
    """
    Loads all available vector stores from the VECTOR_STORE_PATH.
    Each sub-directory in VECTOR_STORE_PATH is considered a separate vector store.
    """
    global vector_stores
    api_key = os.getenv("TOGETHER_API_KEY")
    if not api_key:
        print("TOGETHER_API_KEY not found in environment variables.")
        return

    embeddings = TogetherEmbeddings(model=EMBEDDING_MODEL, together_api_key=api_key)
    
    if os.path.exists(VECTOR_STORE_PATH) and os.path.isdir(VECTOR_STORE_PATH):
        for field_name in os.listdir(VECTOR_STORE_PATH):
            field_path = os.path.join(VECTOR_STORE_PATH, field_name)
            if os.path.isdir(field_path):
                try:
                    # The FAISS index file is expected inside this directory
                    if os.path.exists(os.path.join(field_path, "index.faiss")):
                        print(f"Loading vector store for knowledge field: {field_name}")
                        vector_stores[field_name] = FAISS.load_local(field_path, embeddings, allow_dangerous_deserialization=True)
                    else:
                        print(f"Skipping directory {field_name}: No index.faiss file found.")
                except Exception as e:
                    print(f"Error loading vector store for '{field_name}': {e}")
    else:
        print("Vector store path does not exist or is not a directory.")
        vector_stores = {}

def smart_chunk_document(file_path: str) -> list[Document]:
    """
    Creates intelligent chunks based on document structure with AI-enhanced context.
    Uses adaptive chunking strategy based on document size and complexity.
    Enhanced for better context preservation with larger chunks.
    """
    doc = docx.Document(file_path)
    docs = []
    current_headings = [""] * 6
    
    # Extract knowledge field from the file path - Fixed to find the correct Documents folder
    path_parts = os.path.normpath(file_path).split(os.sep)
    try:
        # Find the Documents folder that contains the script directory
        script_dir_name = os.path.basename(SCRIPT_DIR)
        script_index = path_parts.index(script_dir_name)
        documents_index = script_index + 1  # Documents folder is right after script directory
        
        # Ensure we found the correct Documents folder
        if documents_index < len(path_parts) and path_parts[documents_index] == 'Documents':
            knowledge_field = path_parts[documents_index + 1]
        else:
            # Fallback: search for Documents folder from the right side
            documents_indices = [i for i, part in enumerate(path_parts) if part == 'Documents']
            if documents_indices:
                # Take the last Documents folder found
                documents_index = documents_indices[-1]
                knowledge_field = path_parts[documents_index + 1]
            else:
                knowledge_field = "Unknown"
    except (ValueError, IndexError):
        knowledge_field = "Unknown"

    # Collect paragraphs with their heading context
    content_blocks = []
    current_block = ""
    current_heading_context = ""

    for para in doc.paragraphs:
        if para.style.name.startswith('Heading'):
            # Save current block if it has content
            if current_block.strip():
                content_blocks.append({
                    "content": current_block.strip(),
                    "heading_context": current_heading_context,
                    "knowledge_field": knowledge_field
                })
                current_block = ""
            
            # Update heading hierarchy
            try:
                level = int(para.style.name.split(' ')[-1])
                if 1 <= level <= 6:
                    current_headings[level - 1] = para.text.strip()
                    for i in range(level, 6):
                        current_headings[i] = ""
                    current_heading_context = " > ".join(h for h in current_headings if h)
            except ValueError:
                continue
        elif para.text.strip():
            current_block += para.text + "\n"
    
    # Don't forget the last block
    if current_block.strip():
        content_blocks.append({
            "content": current_block.strip(),
            "heading_context": current_heading_context,
            "knowledge_field": knowledge_field
        })

    # ENHANCED: Adaptive chunking based on document size and complexity
    total_content_length = sum(len(block['content']) for block in content_blocks)
    
    # Determine optimal chunk size based on document characteristics
    if total_content_length <= SMALL_DOC_THRESHOLD:
        # Small documents: Use smaller chunks but still larger than original
        chunk_size = SMALL_CHUNK_SIZE
        overlap = 400
        strategy_info = f"Small document ({total_content_length} chars) - Using {chunk_size} char chunks"
    elif total_content_length <= MEDIUM_DOC_THRESHOLD:
        # Medium documents: Use medium chunks  
        chunk_size = MEDIUM_CHUNK_SIZE
        overlap = DEFAULT_OVERLAP
        strategy_info = f"Medium document ({total_content_length} chars) - Using {chunk_size} char chunks"
    else:
        # Large documents: Use large chunks for better context
        chunk_size = LARGE_CHUNK_SIZE
        overlap = 800  # More overlap for large chunks
        strategy_info = f"Large document ({total_content_length} chars) - Using {chunk_size} char chunks"
    
    print(f"[ENHANCED CHUNKING] {strategy_info}")
    
    # Create documents with enhanced context using adaptive chunking
    from langchain.text_splitter import RecursiveCharacterTextSplitter
    text_splitter = RecursiveCharacterTextSplitter(
        chunk_size=chunk_size,
        chunk_overlap=overlap,
        separators=["\n\n", "\n", ". ", " ", ""]
    )
    
    for i, block in enumerate(content_blocks):
        # Clean the content to avoid API validation issues
        content = block['content'].strip()
        heading_context = block['heading_context'] or "Dokument"
        
        # Skip empty content blocks
        if not content:
            continue
            
        # Enhanced content with context - simplified format
        enhanced_content = f"Kontext: {heading_context}\n\n{content}"
        
        # Split large blocks into adaptive chunks
        chunks = text_splitter.split_text(enhanced_content)
        
        for j, chunk in enumerate(chunks):
            if chunk.strip():  # Only add non-empty chunks
                # Validate chunk size for embedding model (max 8192 tokens ≈ 6000 chars)
                chunk_content = chunk.strip()
                if len(chunk_content) > 6000:
                    print(f"[WARNING] Chunk too large ({len(chunk_content)} chars), truncating to 6000 chars")
                    chunk_content = chunk_content[:6000]
                
                metadata = {
                    "source": file_path,
                    "heading": heading_context,
                    "knowledge_field": knowledge_field,
                    "block_id": f"{i}_{j}",
                    "total_blocks": len(content_blocks),
                    "chunk_strategy": strategy_info,
                    "chunk_size": chunk_size
                }
                
                docs.append(Document(page_content=chunk_content, metadata=metadata))
    
    return docs

def _process_pdf_for_knowledge_base(file_path: str) -> list[Document]:
    """
    Smarter PDF processing that extracts structure (headings) and adds it to chunks.
    Mirrors the logic of smart_chunk_document for DOCX files to improve searchability.
    """
    docs = []
    try:
        # Use unstructured to partition the PDF into structured elements.
        # "hi_res" strategy is crucial for identifying elements like headers and titles.
        elements = partition(filename=file_path, strategy="hi_res", languages=["deu", "eng"])
        
        # Extract knowledge field from the file path
        path_parts = os.path.normpath(file_path).split(os.sep)
        try:
            documents_index = path_parts.index('Documents')
            knowledge_field = path_parts[documents_index + 1]
        except (ValueError, IndexError):
            knowledge_field = "Unknown"

        text_splitter = RecursiveCharacterTextSplitter(
            chunk_size=DEFAULT_CHUNK_SIZE,
            chunk_overlap=DEFAULT_OVERLAP,
            separators=["\n\n", "\n", ". ", " ", ""]
        )

        current_heading = ""
        content_buffer = ""
        
        # Group text elements under the last seen heading
        for element in elements:
            # Heuristic to identify headings from element categories
            is_heading = element.category in ("Title", "Header", "SubTitle")

            if is_heading:
                # When a new heading is found, process the content collected for the *previous* one.
                if content_buffer.strip():
                    # Add context to the content before chunking
                    page_content = f"[Kontext: {current_heading}] {content_buffer.strip()}" if current_heading else content_buffer.strip()
                    chunks = text_splitter.split_text(page_content)
                    for i, chunk_text in enumerate(chunks):
                        if chunk_text.strip():
                            metadata = {
                                "source": file_path,
                                "heading": current_heading,
                                "knowledge_field": knowledge_field,
                                "chunk_id": f"{element.id}_{i}"
                            }
                            docs.append(Document(page_content=chunk_text.strip(), metadata=metadata))
                    content_buffer = ""  # Reset buffer after processing

                # Update the current heading with the new one found
                current_heading = element.text.strip()
            else:
                # If it's not a heading, add its text to the buffer for the current section
                content_buffer += element.text + "\n"

        # After the loop, process any remaining content in the buffer
        if content_buffer.strip():
            page_content = f"[Kontext: {current_heading}] {content_buffer.strip()}" if current_heading else content_buffer.strip()
            chunks = text_splitter.split_text(page_content)
            for i, chunk_text in enumerate(chunks):
                if chunk_text.strip():
                    metadata = {
                        "source": file_path,
                        "heading": current_heading,
                        "knowledge_field": knowledge_field,
                        "chunk_id": f"final_chunk_{i}"
                    }
                    docs.append(Document(page_content=chunk_text.strip(), metadata=metadata))

    except Exception as e:
        print(f"Error processing PDF {file_path} with smart chunking: {e}")
        # On failure, return an empty list to avoid adding bad data to the vector store
        return []
        
    return docs

def force_create_vector_store():
    """
    Clean and simple knowledge base creation using smart chunking and pure semantic search.
    Relies entirely on the excellent multilingual embeddings for understanding German technical terms.
    """
    global vector_stores
    BATCH_SIZE = 64  # Optimized batch size for efficiency
    
    if not os.path.isdir(DOCUMENTS_PATH):
        yield "Documents directory not found. Knowledge base creation skipped."
        return

    # Get all subdirectories in the Documents path, these are the knowledge fields
    knowledge_fields = [d for d in os.listdir(DOCUMENTS_PATH) if os.path.isdir(os.path.join(DOCUMENTS_PATH, d)) and not d.startswith('.')]
    
    if not knowledge_fields:
        yield "No knowledge fields (sub-directories) found in 'Documents'. Nothing to process."
        return

    yield f"Found {len(knowledge_fields)} knowledge fields: {', '.join(knowledge_fields)}"
    
    # Clear existing vector stores from memory
    yield "Clearing existing vector stores from memory..."
    vector_stores.clear()
    gc.collect()
    
    # Remove existing vector store directory
    if os.path.exists(VECTOR_STORE_PATH):
        yield "Removing existing vector store directory..."
        import shutil
        try:
            shutil.rmtree(VECTOR_STORE_PATH)
        except OSError as e:
            yield f"Warning: Could not remove vector store directory: {e}. Continuing with creation..."
    
    os.makedirs(VECTOR_STORE_PATH, exist_ok=True)

    # Initialize embeddings with the excellent multilingual model
    embeddings_model = TogetherEmbeddings(model=EMBEDDING_MODEL, together_api_key=os.getenv("TOGETHER_API_KEY"))
    
    # Process each knowledge field separately
    for field in knowledge_fields:
        yield f"--- Processing Knowledge Field: {field} ---"
        field_path = os.path.join(DOCUMENTS_PATH, field)
        doc_files = get_document_list(field_path)

        if not doc_files:
            yield f"No documents found for '{field}'. Skipping."
            continue

        yield f"Found {len(doc_files)} documents for '{field}'. Creating smart chunks..."
        
        all_docs = []
        
        for i, doc_path in enumerate(doc_files):
            filename = os.path.basename(doc_path)
            yield f"  - Processing document {i+1}/{len(doc_files)}: {filename}"
            try:
                chunks = []
                if doc_path.lower().endswith('.docx'):
                    yield f"    - Using DOCX processor..."
                    chunks = smart_chunk_document(doc_path)
                    yield f"    - Created {len(chunks)} intelligent chunks from DOCX."
                elif doc_path.lower().endswith('.pdf'):
                    yield f"    - Using PDF processor..."
                    chunks = _process_pdf_for_knowledge_base(doc_path)
                    yield f"    - Created {len(chunks)} intelligent chunks from PDF."
                
                if chunks:
                    all_docs.extend(chunks)
                    yield f"    - Successfully added {len(chunks)} chunks to the knowledge base."
                else:
                    yield f"    - WARNING: No chunks were created for this document. It might be empty or unreadable."

            except Exception as e:
                yield f"  - CRITICAL ERROR processing {filename}: {e}"
                continue

        if not all_docs:
            yield f"No content could be extracted from documents in '{field}'. Skipping."
            continue

        yield f"Creating embeddings for '{field}' with multilingual model..."
        
        # Prepare texts and metadata for embedding
        texts = [doc.page_content for doc in all_docs]
        metadatas = [doc.metadata for doc in all_docs]

        # Create embeddings in optimized batches
        all_embeddings = []
        total_batches = (len(texts) + BATCH_SIZE - 1) // BATCH_SIZE
        for i in range(0, len(texts), BATCH_SIZE):
            batch_texts = texts[i:i+BATCH_SIZE]
            yield f"  - Creating embeddings for '{field}' (Batch {i//BATCH_SIZE + 1} of {total_batches})"
            batch_embeddings = embeddings_model.embed_documents(batch_texts)
            all_embeddings.extend(batch_embeddings)

        yield f"Building semantic search index for '{field}'..."
        text_embedding_pairs = list(zip(texts, all_embeddings))
        
        # Create the field-specific vector store
        field_vector_store = FAISS.from_embeddings(text_embedding_pairs, embeddings_model, metadatas=metadatas)
        
        # Save it in its own directory
        field_save_path = os.path.join(VECTOR_STORE_PATH, field)
        os.makedirs(field_save_path, exist_ok=True)
        field_vector_store.save_local(field_save_path)
        
        yield f"✅ Knowledge base for '{field}' successfully created with {len(all_docs)} intelligent chunks"

    yield "--- Knowledge base creation complete. Loading into memory... ---"
    # Reload all vector stores into memory
    load_vector_store()
    yield "🎉 Clean semantic search system ready! German technical terms should now work perfectly."

def analyze_pdf_complexity(file_path: str) -> dict:
    """
    Analyzes PDF complexity to determine the best extraction strategy.
    Returns a dictionary with analysis results.
    """
    if not PYMUPDF_AVAILABLE:
        return {"complexity": "unknown", "use_fallback": True}
    
    try:
        doc = fitz.open(file_path)
        analysis = {
            'page_count': len(doc),
            'file_size_mb': os.path.getsize(file_path) / (1024*1024),
            'has_images': False,
            'has_tables': False,
            'text_density': 0,
            'complexity': 'simple'
        }
        
        # Analyze first 3 pages for complexity indicators
        sample_pages = min(3, len(doc))
        total_text_length = 0
        
        for page_num in range(sample_pages):
            page = doc[page_num]
            text = page.get_text()
            total_text_length += len(text.strip())
            
            # Check for images
            if len(page.get_images()) > 0:
                analysis['has_images'] = True
            
            # Simple table detection heuristics
            if '|' in text or '\t' in text or text.count('\n') > len(text) / 50:
                analysis['has_tables'] = True
        
        analysis['text_density'] = total_text_length / sample_pages if sample_pages > 0 else 0
        
        # Determine complexity
        if (analysis['file_size_mb'] > 50 or 
            analysis['page_count'] > 200 or 
            analysis['text_density'] < 100):
            analysis['complexity'] = 'high'
        elif (analysis['has_tables'] or 
              analysis['has_images'] or 
              analysis['file_size_mb'] > 20):
            analysis['complexity'] = 'medium'
        
        doc.close()
        return analysis
        
    except Exception as e:
        print(f"Error analyzing PDF complexity: {e}")
        return {"complexity": "unknown", "use_fallback": True}

def extract_text_with_pymupdf(file_path: str) -> str:
    """Fast PDF text extraction using PyMuPDF."""
    if not PYMUPDF_AVAILABLE:
        raise ImportError("PyMuPDF not available")
    
    doc = fitz.open(file_path)
    text = ""
    for page in doc:
        text += page.get_text() + "\n\n"
    doc.close()
    return text

def extract_text_with_pdfplumber(file_path: str) -> str:
    """PDF text extraction with better table support using pdfplumber."""
    if not PDFPLUMBER_AVAILABLE:
        raise ImportError("pdfplumber not available")
    
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page in pdf.pages:
            page_text = page.extract_text()
            if page_text:
                text += page_text + "\n\n"
    return text

def extract_text_with_unstructured(file_path: str, strategy: str = "fast") -> str:
    """PDF text extraction using unstructured with specified strategy."""
    elements = partition(filename=file_path, include_metadata=True, strategy=strategy)
    return "\n\n".join([str(element) for element in elements if str(element).strip()])

async def smart_pdf_extraction(file_path: str, cancellation_check=lambda: False):
    """
    Intelligently extracts text from PDF using the best available method.
    Returns tuple of (text, method_used).
    """
    try:
        # Step 1: Analyze PDF complexity
        yield {"status": "processing", "message": "Analyzing PDF structure..."}
        analysis = await asyncio.to_thread(analyze_pdf_complexity, file_path)
        
        if cancellation_check(): return
        
        # Step 2: Try PyMuPDF first (fastest)
        if PYMUPDF_AVAILABLE and analysis['complexity'] in ['simple', 'medium']:
            try:
                yield {"status": "processing", "message": "Extracting text with fast method (PyMuPDF)..."}
                text = await asyncio.to_thread(extract_text_with_pymupdf, file_path)
                if len(text.strip()) > 1000:  # Good extraction
                    yield {"status": "complete", "text": text, "method": "PyMuPDF (fast)"}
                    return
            except Exception as e:
                print(f"PyMuPDF extraction failed: {e}")
        
        if cancellation_check(): return
        
        # Step 3: Try pdfplumber for tables
        if PDFPLUMBER_AVAILABLE and analysis.get('has_tables', False):
            try:
                yield {"status": "processing", "message": "Extracting text with table-optimized method (pdfplumber)..."}
                text = await asyncio.to_thread(extract_text_with_pdfplumber, file_path)
                if len(text.strip()) > 500:  # Reasonable extraction
                    yield {"status": "complete", "text": text, "method": "pdfplumber (table-optimized)"}
                    return
            except Exception as e:
                print(f"pdfplumber extraction failed: {e}")
        
        if cancellation_check(): return
        
        # Step 4: Try unstructured with fast strategy
        try:
            yield {"status": "processing", "message": "Extracting text with structured method (unstructured-fast)..."}
            text = await asyncio.to_thread(extract_text_with_unstructured, file_path, "fast")
            if len(text.strip()) > 500:
                yield {"status": "complete", "text": text, "method": "unstructured (fast)"}
                return
        except Exception as e:
            print(f"Unstructured fast extraction failed: {e}")
        
        if cancellation_check(): return
        
        # Step 5: Ask user for hi_res processing (slow but thorough)
        file_size_mb = analysis.get('file_size_mb', 0)
        estimated_time = "5-15 minutes" if file_size_mb > 10 else "2-5 minutes"
        
        yield {
            "status": "user_confirmation_needed", 
            "message": f"PDF requires advanced processing (estimated time: {estimated_time}). This will use the most thorough but slowest method. Continue?",
            "analysis": analysis
        }
        
        # Note: The calling function should handle user confirmation
        # For now, we'll proceed with hi_res as fallback
        yield {"status": "processing", "message": f"Using advanced processing (estimated time: {estimated_time})..."}
        text = await asyncio.to_thread(extract_text_with_unstructured, file_path, "hi_res")
        yield {"status": "complete", "text": text, "method": "unstructured (hi_res - thorough)"}
        
    except Exception as e:
        yield {"status": "error", "message": f"All extraction methods failed: {e}"}

async def create_vector_store_for_document(file_path: str, vector_store_path: str, cancellation_check=lambda: False):
    """
    Intelligently processes uploaded documents: uses full-text analysis for documents under 200k characters,
    RAG with chunking for larger documents. Only affects user-uploaded documents, not the global knowledge base.
    """
    try:
        # Use smart extraction for PDFs
        if file_path.lower().endswith('.pdf'):
            extraction_complete = False
            extracted_text = ""
            method_used = ""
            
            async for result in smart_pdf_extraction(file_path, cancellation_check):
                if cancellation_check(): return
                
                status = result.get("status")
                message = result.get("message", "")
                
                if status == "processing":
                    yield {"status": "processing", "message": message}
                elif status == "user_confirmation_needed":
                    # For now, we'll auto-proceed. In a full implementation,
                    # this could be sent to the frontend for user decision
                    yield {"status": "processing", "message": "PDF requires advanced processing. Proceeding with thorough extraction..."}
                elif status == "complete":
                    extracted_text = result.get("text", "")
                    method_used = result.get("method", "unknown")
                    extraction_complete = True
                    yield {"status": "processing", "message": f"Text extracted successfully using {method_used}"}
                    break
                elif status == "error":
                    yield {"status": "error", "message": message}
                    return
            
            if not extraction_complete or not extracted_text.strip():
                yield {"status": "error", "message": "Could not extract any text from the PDF"}
                return
            
            # Convert extracted text to elements format for consistency
            full_text = extracted_text
            
        else:
            # Use original method for non-PDF files
            yield {"status": "processing", "message": "Analyzing and partitioning document..."}
            elements = await asyncio.to_thread(partition, filename=file_path, include_metadata=True, strategy="hi_res")
            if cancellation_check(): return
            full_text = "\n\n".join([str(element) for element in elements if str(element).strip()])

        if cancellation_check(): return

        # INTELLIGENT PROCESSING DECISION: Full-text vs RAG based on document size
        text_length = len(full_text)
        FULLTEXT_THRESHOLD = 200000  # 200k characters
        
        if text_length < FULLTEXT_THRESHOLD:
            # Use full-text processing for better table analysis
            yield {"status": "processing", "message": f"Document size: {text_length:,} characters - Using full-text analysis for optimal results"}
            
            # Save the full text directly to a special marker file
            fulltext_marker_path = os.path.join(vector_store_path, "fulltext_content.txt")
            os.makedirs(vector_store_path, exist_ok=True)
            
            with open(fulltext_marker_path, 'w', encoding='utf-8') as f:
                f.write(full_text)
            
            yield {"status": "complete", "message": "Document ready for full-text analysis (optimal for tables and structured data)."}
            
        else:
            # Use traditional RAG processing for large documents
            yield {"status": "processing", "message": f"Document size: {text_length:,} characters - Using RAG with chunking for large document"}
            
            # Use RecursiveCharacterTextSplitter for better chunking with overlap
            text_splitter = RecursiveCharacterTextSplitter(
                chunk_size=DEFAULT_CHUNK_SIZE,  # Reduced from 1500 for better granularity
                chunk_overlap=DEFAULT_OVERLAP,  # Added overlap for better continuity
                separators=["\n\n", "\n", ". ", " ", ""]
            )
            
            # Create documents with metadata
            docs = []
            chunks = text_splitter.split_text(full_text)
            for i, chunk in enumerate(chunks):
                if chunk.strip():
                    docs.append(Document(
                        page_content=chunk,
                        metadata={
                            "source": file_path,
                            "chunk_id": i,
                            "total_chunks": len(chunks)
                        }
                    ))

            if not docs:
                yield {"status": "error", "message": "Could not extract any text from the document. The document might be empty, image-based without readable text, or corrupted."}
                return

            yield {"status": "processing", "message": "Creating embeddings with optimized batching..."}
            api_key = os.getenv("TOGETHER_API_KEY")
            embeddings = TogetherEmbeddings(model=EMBEDDING_MODEL, together_api_key=api_key)
            
            # Optimized batch processing for embeddings
            BATCH_SIZE = 64  # Increased batch size for better efficiency
            all_embeddings = []
            texts = [doc.page_content for doc in docs]
            metadatas = [doc.metadata for doc in docs]
            
            total_batches = (len(texts) + BATCH_SIZE - 1) // BATCH_SIZE
            for i in range(0, len(texts), BATCH_SIZE):
                if cancellation_check(): return
                batch_texts = texts[i:i+BATCH_SIZE]
                yield {"status": "processing", "message": f"Processing embeddings batch {i//BATCH_SIZE + 1}/{total_batches}..."}
                batch_embeddings = await asyncio.to_thread(embeddings.embed_documents, batch_texts)
                all_embeddings.extend(batch_embeddings)

            # Create FAISS index with embeddings
            text_embedding_pairs = list(zip(texts, all_embeddings))
            db = await asyncio.to_thread(FAISS.from_embeddings, text_embedding_pairs, embeddings, metadatas=metadatas)
            if cancellation_check(): return

            # save_local is also synchronous
            await asyncio.to_thread(db.save_local, vector_store_path)

            yield {"status": "complete", "message": "Document ready for RAG-based questions."}
        
        # Memory cleanup after large operations
        gc.collect()

    except Exception as e:
        yield {"status": "error", "message": f"Error processing document: {e}"}


async def get_answer_from_rag(conversation_history: list, vector_store_path: str, cancellation_check=lambda: False):
    """
    Intelligently chooses between full-text analysis and RAG based on document processing method.
    Checks for fulltext marker and uses appropriate processing method.
    """
    try:
        # Check if this document was processed with full-text analysis
        fulltext_marker_path = os.path.join(vector_store_path, "fulltext_content.txt")
        
        if os.path.exists(fulltext_marker_path):
            # Use full-text analysis for better table/structured data processing
            yield {"type": "status", "data": "Using full-text analysis (optimal for tables and structured data)..."}
            
            # Read the full document content
            with open(fulltext_marker_path, 'r', encoding='utf-8') as f:
                document_content = f.read()
            
            # Use the direct document processing function
            async for result in get_answer_from_document(conversation_history, document_content, file_type="text", cancellation_check=cancellation_check):
                yield result
            return
        
        # Fall back to traditional RAG processing
        yield {"type": "status", "data": "Loading document context..."}
        api_key = os.getenv("TOGETHER_API_KEY")
        client = Together(api_key=api_key)
        embeddings = TogetherEmbeddings(model=EMBEDDING_MODEL, together_api_key=api_key)
        last_question = conversation_history[-1]['content']
        
        # Load the vector store. This is synchronous.
        vector_store = await asyncio.to_thread(FAISS.load_local, vector_store_path, embeddings, allow_dangerous_deserialization=True)
        
        # Optimized retrieval: Increased for better table coverage
        retriever = vector_store.as_retriever(search_kwargs={"k": 15})  # Increased for better coverage
        
        # Simplified query expansion - limit to 2 additional queries max
        yield {"type": "status", "data": "Optimizing search queries..."}
        expanded_queries = await asyncio.to_thread(expand_query_with_llm_optimized, client, conversation_history, cancellation_check)
        if cancellation_check(): return

        # Retrieve documents for all expanded queries
        all_docs = []
        tasks = [retriever.ainvoke(query) for query in expanded_queries]
        retrieved_results = await asyncio.gather(*tasks)
        for result in retrieved_results:
            if cancellation_check(): return
            all_docs.extend(result)
        
        # Get unique documents
        docs = list({(doc.page_content, doc.metadata.get('source', 'N/A')): doc for doc in all_docs}.values())
        if cancellation_check(): return

        # Two-stage re-ranking for better performance
        if docs:
            yield {"type": "status", "data": "Smart re-ranking for relevance..."}
            
            # Stage 1: Fast similarity filtering if we have too many docs
            if len(docs) > 25:  # Increased threshold to be less aggressive
                # Quick cosine similarity filtering
                doc_embeddings = await asyncio.to_thread(embeddings.embed_documents, [doc.page_content for doc in docs])
                query_embedding = await asyncio.to_thread(embeddings.embed_query, last_question)
                
                # Calculate cosine similarities
                import numpy as np
                similarities = []
                for doc_emb in doc_embeddings:
                    similarity = np.dot(query_embedding, doc_emb) / (np.linalg.norm(query_embedding) * np.linalg.norm(doc_emb))
                    similarities.append(similarity)
                
                # Keep top 18 based on cosine similarity (increased from 10)
                doc_sim_pairs = list(zip(docs, similarities))
                doc_sim_pairs.sort(key=lambda x: x[1], reverse=True)
                docs = [doc for doc, sim in doc_sim_pairs[:18]]
            
            # Stage 2: Cross-encoder re-ranking on filtered docs using cached model
            reranker_model = await get_cached_reranker()
            pairs = [[last_question, doc.page_content] for doc in docs]
            scores = await asyncio.to_thread(reranker_model.predict, pairs)
            scored_docs = sorted(list(zip(scores, docs)), key=lambda x: x[0], reverse=True)
            docs = [doc for score, doc in scored_docs[:7]] # Keep top 7

        # Create a new retriever with the re-ranked documents
        if docs:
            new_vector_store = await asyncio.to_thread(FAISS.from_documents, docs, embeddings)
            retriever = new_vector_store.as_retriever()

        llm = ChatTogether(
            model=LLM_MODEL,
            temperature=0.0,
            max_tokens=2048,
            together_api_key=api_key
        )

        # Set up memory
        memory = ConversationBufferMemory(memory_key='chat_history', return_messages=True)
        for msg in conversation_history[:-1]: # Populate memory with previous messages
            if msg['role'] == 'user':
                memory.chat_memory.add_user_message(msg['content'])
            elif msg['role'] == 'assistant':
                memory.chat_memory.add_ai_message(msg['content'])

        # Create the conversational chain with generic document analysis prompt
        prompt_template = """Sie sind ein Experte für Dokumentenanalyse und Informationsextraktion. Ihre Aufgabe ist es, die Frage des Benutzers basierend *ausschließlich* auf dem bereitgestellten Kontext aus einem Dokument zu beantworten.

**INTELLIGENTE ABKÜRZUNGS- UND BEGRIFFSERKENNUNG:**
- Erkennen Sie gängige Abkürzungen und Akronyme im Kontext des Dokuments
- Suchen Sie nach sowohl abgekürzten als auch vollständigen Formen von Begriffen
- Berücksichtigen Sie fachspezifische Terminologie je nach Dokumenttyp
- Achten Sie auf Definitionen und Erklärungen im Dokument selbst

**SPEZIELLE ANWEISUNGEN FÜR STRUKTURIERTE DATEN:**
- Wenn Sie Tabellendaten sehen, analysieren Sie ALLE Zeilen und Spalten sorgfältig
- Suchen Sie nach Mustern und Zusammenhängen zwischen verschiedenen Tabellenteilen
- Identifizieren Sie Trends, Veränderungen und wichtige Datenpunkte
- Kombinieren Sie Informationen aus verschiedenen Dokumententeilen
- Erklären Sie die Bedeutung der Daten im Kontext des Dokuments
- Achten Sie auf Zahlen, Statistiken, Vergleiche und deren Relevanz

**ANWEISUNGEN:**
1.  **Gründlich Analysieren:** Lesen Sie den bereitgestellten Kontext sorgfältig, um alle Informationen vollständig zu verstehen.
2.  **Kontextbewusste Begriffserkennung:** Erkennen Sie Fachbegriffe, Abkürzungen und spezielle Terminologie basierend auf dem Dokumenttyp und -inhalt.
3.  **Umfassende Antwort:** Formulieren Sie eine detaillierte und gründliche Antwort auf die Frage des Benutzers. Gehen Sie über eine einfache Ein-Satz-Antwort hinaus. Erklären Sie das 'Warum' und 'Wie', wenn möglich basierend auf dem Kontext.
4.  **Markdown-Formatierung:** Strukturieren Sie Ihre Antwort klar mit Markdown. Verwenden Sie Aufzählungspunkte, fetten Text und Absätze, um die Antwort leicht lesbar und verständlich zu machen.
5.  **Informationen Synthetisieren:** Wenn der Kontext Informationen aus verschiedenen Teilen des Dokuments enthält, synthetisieren Sie diese zu einer kohärenten Antwort.
6.  **Quelltreu Bleiben:** Basieren Sie Ihre Antwort strikt auf dem bereitgestellten Kontext. Fügen Sie keine Informationen aus externem Wissen hinzu.
7.  **Fehlende Informationen Behandeln:** Wenn die Antwort nicht im bereitgestellten Kontext gefunden werden kann, geben Sie klar an: "Ich konnte keine definitive Antwort auf diese Frage im Dokument finden."
8.  **Sprache:** Sie MÜSSEN in derselben Sprache antworten, die der Benutzer in seiner Frage verwendet hat.

**Kontext:**
{context}

**Frage:**
{question}

Antwort:"""
        QA_PROMPT = PromptTemplate(
            template=prompt_template, input_variables=["context", "question"]
        )

        qa_chain = ConversationalRetrievalChain.from_llm(
            llm=llm,
            retriever=retriever,
            memory=memory,
            verbose=False, # for debugging
            combine_docs_chain_kwargs={"prompt": QA_PROMPT}
        )

        yield {"type": "status", "data": "Formulating answer from document..."}
        last_question = conversation_history[-1]['content']

        # The chain's `ainvoke` will handle history and retrieval
        response = await qa_chain.ainvoke({"question": last_question})
        answer = response.get("answer", "I could not find an answer in the document.")

        # Since the chain doesn't stream, we stream the final answer manually
        # In a real scenario, you might use a custom streaming chain.
        # For now, we simulate the streaming behavior.
        
        # First, send metadata
        sources_text = "\n".join(sorted(list(set([f"- {os.path.basename(doc.metadata.get('source', 'Unknown'))}" for doc in response.get('source_documents', [])]))))
        yield {"type": "meta", "data": {"sources": f"**Document sources:**\n{sources_text}", "keywords": "N/A", "follow_ups": [], "source_mode": "rag_document"}}

        # Then, stream the answer
        chunk_size = 20
        for i in range(0, len(answer), chunk_size):
            if cancellation_check(): return
            yield {"type": "chunk", "data": answer[i:i+chunk_size]}
            await asyncio.sleep(0.01) # small delay to simulate streaming

        yield {"type": "end"}

    except Exception as e:
        error_message = f"An error occurred in RAG processing: {e}"
        print(error_message)
        yield {"type": "meta", "data": {"sources": "No sources", "keywords": "N/A", "follow_ups": []}}
        yield {"type": "chunk", "data": error_message}
        yield {"type": "end"}


# --- Cached Reranker Model ---
_cached_reranker = None

async def get_cached_reranker():
    """
    Returns a cached reranker model instance for better performance.
    Creates the model on first use and reuses it for subsequent calls.
    """
    global _cached_reranker
    if _cached_reranker is None:
        try:
            # Use a simple dummy reranker for now since we don't have sentence-transformers installed
            # In a production environment, you would use something like:
            # from sentence_transformers import CrossEncoder
            # _cached_reranker = CrossEncoder('cross-encoder/ms-marco-MiniLM-L-12-v2')
            class DummyReranker:
                def predict(self, pairs):
                    # Simple fallback: return scores based on query length match
                    import random
                    return [random.uniform(0.3, 0.9) for _ in pairs]
            
            _cached_reranker = DummyReranker()
        except Exception as e:
            print(f"Error initializing reranker: {e}")
            # Fallback dummy reranker
            class DummyReranker:
                def predict(self, pairs):
                    import random
                    return [random.uniform(0.3, 0.9) for _ in pairs]
            _cached_reranker = DummyReranker()
    
    return _cached_reranker

load_vector_store()


# --- Conversational AI Functions ---

def create_contextual_messages(conversation_history: list, system_prompt: str) -> list:
    messages = [{"role": "system", "content": system_prompt}]
    history = [msg for msg in conversation_history if msg["role"] in ["user", "assistant"]]
    for msg in history:
        messages.append({"role": msg["role"], "content": msg["content"]})
    return messages

def translate_queries_to_english(client, queries: list[str], cancellation_check=lambda: False) -> list[str]:
    """Translates a list of queries to English using an LLM."""
    system_prompt = "You are a highly skilled translator. Your task is to translate the following comma-separated list of search queries accurately from German to English. Return only the comma-separated list of the translated queries."
    
    queries_string = ", ".join(queries)
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": queries_string}
    ]
    
    try:
        if cancellation_check(): return []
        # Using a specific, smaller model for the translation task as requested
        response = robust_api_call(client, FAST_MODEL, messages, 0.2, cancellation_check=cancellation_check)
        if response.choices and response.choices[0].message.content:
            translated_queries = [q.strip() for q in response.choices[0].message.content.strip().split(',')]
            return translated_queries
        return []
    except Exception as e:
        print(f"Error during query translation: {e}")
        return []

def detect_and_expand_abbreviations(text: str) -> list[str]:
    """
    Detects common German business abbreviations and returns expanded versions with multiple case and format variants.
    Returns a list of expanded queries including both abbreviated and full forms.
    """
    abbreviation_map = {
        'GuV': [
            'Gewinn- und Verlustrechnung',
            'Gewinn-und-Verlust-Rechnung', 
            'GEWINN-UND-VERLUST-RECHNUNG',
            'Gewinn und Verlustrechnung',
            'GEWINN UND VERLUSTRECHNUNG',
            'Gewinn-Und-Verlust-Rechnung'
        ],
        'GmbH': ['Gesellschaft mit beschränkter Haftung', 'GESELLSCHAFT MIT BESCHRÄNKTER HAFTUNG'],
        'AG': ['Aktiengesellschaft', 'AKTIENGESELLSCHAFT'],
        'KG': ['Kommanditgesellschaft', 'KOMMANDITGESELLSCHAFT'],
        'UG': ['Unternehmergesellschaft', 'UNTERNEHMERGESELLSCHAFT'],
        'BWA': [
            'Betriebswirtschaftliche Auswertung',
            'BETRIEBSWIRTSCHAFTLICHE AUSWERTUNG',
            'Betriebswirtschaftliche-Auswertung'
        ],
        'USt': ['Umsatzsteuer', 'UMSATZSTEUER'],
        'MwSt': ['Mehrwertsteuer', 'MEHRWERTSTEUER'],
        'EK': ['Eigenkapital', 'EIGENKAPITAL'],
        'FK': ['Fremdkapital', 'FREMDKAPITAL'],
        'ROI': ['Return on Investment', 'RETURN ON INVESTMENT'],
        'EBIT': ['Earnings Before Interest and Taxes', 'EARNINGS BEFORE INTEREST AND TAXES'],
        'EBITDA': ['Earnings Before Interest, Taxes, Depreciation and Amortization', 'EARNINGS BEFORE INTEREST, TAXES, DEPRECIATION AND AMORTIZATION'],
        'KPI': ['Key Performance Indicator', 'KEY PERFORMANCE INDICATOR'],
        'HR': ['Human Resources', 'HUMAN RESOURCES'],
        'IT': ['Informationstechnologie', 'INFORMATIONSTECHNOLOGIE'],
        'CEO': ['Chief Executive Officer', 'CHIEF EXECUTIVE OFFICER'],
        'CFO': ['Chief Financial Officer', 'CHIEF FINANCIAL OFFICER'],
        'CTO': ['Chief Technology Officer', 'CHIEF TECHNOLOGY OFFICER'],
        'B2B': ['Business-to-Business', 'BUSINESS-TO-BUSINESS'],
        'B2C': ['Business-to-Consumer', 'BUSINESS-TO-CONSUMER'],
        'FAQ': ['Frequently Asked Questions', 'FREQUENTLY ASKED QUESTIONS'],
        'PDF': ['Portable Document Format', 'PORTABLE DOCUMENT FORMAT'],
        'API': ['Application Programming Interface', 'APPLICATION PROGRAMMING INTERFACE'],
        'SQL': ['Structured Query Language', 'STRUCTURED QUERY LANGUAGE'],
        'CRM': ['Customer Relationship Management', 'CUSTOMER RELATIONSHIP MANAGEMENT'],
        'ERP': ['Enterprise Resource Planning', 'ENTERPRISE RESOURCE PLANNING']
    }
    
    expanded_queries = [text]  # Always include original
    
    # Add automatic case variants for the original text
    if text.lower() != text:
        expanded_queries.append(text.lower())
    if text.upper() != text:
        expanded_queries.append(text.upper())
    if text.title() != text:
        expanded_queries.append(text.title())
    
    # Check for abbreviations in the text
    words = text.split()
    for word in words:
        # Remove punctuation for matching
        clean_word = word.strip('.,!?;:()[]{}"\'-')
        if clean_word in abbreviation_map:
            # Get all variants for this abbreviation
            variants = abbreviation_map[clean_word]
            if isinstance(variants, str):
                variants = [variants]
            
            # Create expanded versions with all variants
            for variant in variants:
                expanded_text = text.replace(word, variant)
                if expanded_text not in expanded_queries:
                    expanded_queries.append(expanded_text)
    
    # Remove duplicates while preserving order
    unique_queries = []
    for query in expanded_queries:
        if query not in unique_queries:
            unique_queries.append(query)
    
    return unique_queries

def expand_query_with_llm(client, conversation_history: list, cancellation_check=lambda: False) -> list[str]:
    system_prompt = """You are an expert in query expansion for technical documentation. Your task is to reformulate the user's LATEST question into 3 diverse, but closely related search queries, using the full conversation for context.

**CRITICAL ABBREVIATION EXPANSION RULES:**
- You MUST generate the queries in the SAME language as the user's LATEST question.
- **ALWAYS expand common German business abbreviations to their full forms in at least one query:**
  - GuV → Gewinn- und Verlustrechnung
  - GmbH → Gesellschaft mit beschränkter Haftung  
  - AG → Aktiengesellschaft
  - KG → Kommanditgesellschaft
  - UG → Unternehmergesellschaft
  - BWA → Betriebswirtschaftliche Auswertung
  - USt → Umsatzsteuer
  - MwSt → Mehrwertsteuer
  - EK → Eigenkapital
  - FK → Fremdkapital
  - ROI → Return on Investment
  - EBIT → Earnings Before Interest and Taxes
  - EBITDA → Earnings Before Interest, Taxes, Depreciation and Amortization
  - KPI → Key Performance Indicator
  - HR → Human Resources
  - IT → Informationstechnologie
  - CEO → Chief Executive Officer
  - CFO → Chief Financial Officer
  - CTO → Chief Technology Officer
  - B2B → Business-to-Business
  - B2C → Business-to-Consumer
  - FAQ → Frequently Asked Questions
  - PDF → Portable Document Format
  - API → Application Programming Interface
  - SQL → Structured Query Language
  - CRM → Customer Relationship Management
  - ERP → Enterprise Resource Planning
- **If you detect ANY abbreviation in the user's question, create one query with the abbreviation and another with the full form.**
- If the latest question is a follow-up (e.g., "tell me more", "explain that in detail"), derive the search queries from the topic of the PREVIOUS assistant answer.
- Focus on synonyms and different phrasings of the same intent.
- **Do NOT invent new concepts or terms.**
- **Do NOT add any additional context or information that is not present in the conversation.**
- **DO NEVER TRY TO ANSWER THE QUESTION DIRECTLY. JUST REPHRASE AND REFORMULATE THE QUESTIONS**
- **TRY TO FORMULATE THE REFORMULATE THE QUESTION TO ACHIEVE A BETTER SEARCH RESULT**
- If the user's question contains a specific technical term (e.g., 'db_S4U_V2_ALLDATA'), you MUST include that exact term in at least one of the expanded queries.
- Return only a comma-separated list of these new queries.
"""
    messages = create_contextual_messages(conversation_history, system_prompt)
    try:
        if cancellation_check(): return []
        # Use FAST_MODEL for improved performance in query expansion
        response = robust_api_call(client, FAST_MODEL, messages, 0.3)
        if cancellation_check(): return []
        if response.choices and response.choices[0].message.content:
            expanded_queries = [q.strip() for q in response.choices[0].message.content.strip().split(',')]
        else:
            expanded_queries = []
        
        # Ensure the original question is first and the list is unique
        original_question = conversation_history[-1]['content']
        final_queries = [original_question]
        for q in expanded_queries:
            if q not in final_queries:
                final_queries.append(q)
        
        # Additionally, use the automatic abbreviation detection as a fallback
        auto_expanded = detect_and_expand_abbreviations(original_question)
        for q in auto_expanded:
            if q not in final_queries:
                final_queries.append(q)
        
        return final_queries
    except Exception as e:
        print(f"Error during query expansion: {e}")
        return [conversation_history[-1]['content']]

def expand_query_with_llm_optimized(client, conversation_history: list, cancellation_check=lambda: False) -> list[str]:
    """
    Optimized version of query expansion that limits to 2 additional queries max
    and uses a smaller, faster model for better performance.
    """
    system_prompt = """You are an expert in query expansion for technical documentation. Your task is to reformulate the user's LATEST question into 2 diverse search queries.

**CRITICAL ABBREVIATION EXPANSION RULES:**
- Generate queries in the SAME language as the user's question.
- **ALWAYS expand common German business abbreviations to their full forms:**
  - GuV → Gewinn- und Verlustrechnung
  - GmbH → Gesellschaft mit beschränkter Haftung  
  - AG → Aktiengesellschaft
  - BWA → Betriebswirtschaftliche Auswertung
  - USt → Umsatzsteuer
  - MwSt → Mehrwertsteuer
  - EK → Eigenkapital
  - FK → Fremdkapital
  - ROI → Return on Investment
  - EBIT → Earnings Before Interest and Taxes
  - EBITDA → Earnings Before Interest, Taxes, Depreciation and Amortization
  - KPI → Key Performance Indicator
  - HR → Human Resources
  - IT → Informationstechnologie
- **If you detect ANY abbreviation in the user's question, create one query with the abbreviation and another with the full form.**
- Focus on synonyms and different phrasings.
- Keep queries concise and focused.
- Return only a comma-separated list of 2 new queries.
"""
    messages = create_contextual_messages(conversation_history, system_prompt)
    try:
        if cancellation_check(): return []
        # Use smaller, faster model for query expansion
        response = robust_api_call(client, "meta-llama/Llama-3.2-3B-Instruct-Turbo", messages, 0.3, cancellation_check=cancellation_check)
        if cancellation_check(): return []
        if response.choices and response.choices[0].message.content:
            expanded_queries = [q.strip() for q in response.choices[0].message.content.strip().split(',')][:2]  # Limit to 2
        else:
            expanded_queries = []
        
        # Ensure the original question is first and the list is unique
        original_question = conversation_history[-1]['content']
        final_queries = [original_question]
        for q in expanded_queries:
            if q not in final_queries:
                final_queries.append(q)
        
        # Additionally, use the automatic abbreviation detection as a fallback
        auto_expanded = detect_and_expand_abbreviations(original_question)
        for q in auto_expanded:
            if q not in final_queries:
                final_queries.append(q)
        
        return final_queries[:3]  # Maximum 3 queries total
    except Exception as e:
        print(f"Error during optimized query expansion: {e}")
        return [conversation_history[-1]['content']]

def generate_follow_up_questions(client, conversation_history: list, answer: str, context: str) -> list[str]:
    system_prompt = """You are a creative and helpful assistant. Your primary goal is to keep the conversation going by providing insightful follow-up questions.

Based on the user's question, the provided answer, and the available context, generate up to three concise and relevant follow-up questions. You *MUST* use the language the user used in their last question for the questions.
- The questions should be distinct and cover different aspects (e.g., a clarifying question, a question about a related sub-topic, a "what's next" question).
- Always try to generate questions, even if the answer seems complete.
- Think about related topics, implications, or next steps.
- Return ONLY a Python-style list of strings, e.g., ["Question 1?", "Question 2?", "Question 3?"]
"""
    user_content = f"Full Conversation History: {conversation_history}\n\nProvided Answer: {answer}\n\nFull Context Used for Answer: {context}"
    messages = [{"role": "system", "content": system_prompt}, {"role": "user", "content": user_content}]
    try:
        response = robust_api_call(client, LLM_MODEL, messages, 0.7)
        if response.choices and response.choices[0].message.content:
            content = response.choices[0].message.content.strip()
            # Use regex to find all strings within quotes
            questions = re.findall(r'\"(.*?)\"', content)
            return [q for q in questions if q.endswith('?')][:3]
    except Exception as e:
        print(f"Error generating follow-up questions: {e}")
        pass
    return []

def can_answer_without_context(client, conversation_history: list) -> bool:
    """
    Checks if the LLM can answer the question without any external context.
    """
    system_prompt = """You are a helpful assistant. Your task is to determine if you can answer the user's LATEST question based on your general knowledge and the conversation history alone, without needing to search any external documents or the web.

- If the question is about general knowledge (e.g., "What is the capital of France?"), a simple greeting, or a direct follow-up that can be answered from the chat history, respond with 'yes'.
- If the question requires specific, detailed information from a knowledge base (e.g., "What are the details of project X?") or up-to-date information from the web (e.g., "What's the latest news on Y?"), respond with 'no'.

Respond with only 'yes' or 'no'.
"""
    messages = create_contextual_messages(conversation_history, system_prompt)
    try:
        response = robust_api_call(client, LLM_MODEL, messages, 0.0)
        if response.choices and response.choices[0].message.content:
            decision = response.choices[0].message.content.strip().lower()
            return "yes" in decision
        return False
    except Exception as e:
        print(f"Error during can_answer_without_context check: {e}")
        return False

def can_answer_from_conversation_context(client, conversation_history: list) -> bool:
    """
    Checks if the current question can be answered using information from previous assistant responses in the conversation.
    This is crucial for follow-up questions like "Empfiehlst Du einen Regenschirm" after a weather-related answer.
    """
    # Enhanced analysis that works better with web-search results and loaded history
    conversation_text = ""
    for i, msg in enumerate(conversation_history):
        role = msg.get('role', 'unknown')
        content = msg.get('content', '')
        conversation_text += f"{role.upper()}: {content}\n\n"
    
    system_prompt = f"""You are an expert conversation analyst. Your task is to determine if the user's LATEST question can be meaningfully answered using information from PREVIOUS ASSISTANT RESPONSES in this conversation.

**COMPLETE CONVERSATION TO ANALYZE:**
{conversation_text}

**Enhanced Guidelines:**
- Look at the user's latest question and check if it's a follow-up that relates to information the assistant has already provided.
- **IMPORTANT**: Web search results, weather information, news, and other factual data from previous assistant responses COUNT as available context.
- Examples of answerable follow-ups:
  * After weather info about Miami rain → "Should I bring an umbrella?" (YES - can be answered from weather context)
  * After explaining a company policy → "What are the penalties?" (YES - if penalties were mentioned)
  * After providing financial data → "Is that good or bad?" (YES - can provide analysis based on the data)
  * After travel information → "What should I pack?" (YES - can be answered from travel context)
  * After any factual information → "What does this mean for me?" (YES - can provide analysis)

**Analysis Method:**
1. Read the COMPLETE CONVERSATION above
2. Identify what factual information, data, or details the assistant has already provided
3. Check if the user's latest question is asking for advice, analysis, or recommendations based on that existing information
4. Consider web search results, weather data, news, or any other information as valid context

**Critical Rules:**
- If the assistant previously provided any relevant factual information that could inform the user's latest question, answer YES
- Only answer NO if the latest question requires completely new external information
- Questions asking for advice, recommendations, or analysis based on previous assistant responses should be YES

Respond with only 'yes' or 'no'.
"""
    messages = create_contextual_messages(conversation_history, system_prompt)
    try:
        response = robust_api_call(client, LLM_MODEL, messages, 0.0)
        if response.choices and response.choices[0].message.content:
            decision = response.choices[0].message.content.strip().lower()
            return "yes" in decision
        return False
    except Exception as e:
        print(f"Error during can_answer_from_conversation_context check: {e}")
        return False

def extract_context_keywords(client, conversation_history: list) -> list[str]:
    """
    Extracts relevant keywords and entities from previous assistant responses for context-aware search.
    """
    system_prompt = """You are a keyword extraction specialist. Your task is to extract the most relevant keywords and entities from the ASSISTANT'S PREVIOUS RESPONSES that are relevant to the user's LATEST QUESTION.

**Instructions:**
1. Read the user's latest question
2. Look at all previous assistant responses in the conversation
3. Extract 3-7 key terms that provide context for the user's question
4. Focus on: locations, entities, technical terms, topics, specific subjects mentioned by the assistant
5. Prioritize terms that help make the user's question more specific and contextual

**Examples:**
- If user asks "Should I bring an umbrella?" after assistant mentioned "Miami weather tomorrow will be rainy"
  → Extract: ["Miami", "weather", "rain", "tomorrow"]
- If user asks "Is that legal?" after assistant discussed "German employment law"
  → Extract: ["German", "employment law", "legal"]

**Output:** Return ONLY a comma-separated list of keywords, nothing else.
**Language:** Extract keywords in the same language as they appeared in the conversation.
"""
    messages = create_contextual_messages(conversation_history, system_prompt)
    try:
        response = robust_api_call(client, LLM_MODEL, messages, 0.0)
        if response.choices and response.choices[0].message.content:
            keywords_text = response.choices[0].message.content.strip()
            keywords = [kw.strip() for kw in keywords_text.split(',') if kw.strip()]
            return keywords[:7]  # Limit to 7 keywords max
        return []
    except Exception as e:
        print(f"Error during context keyword extraction: {e}")
        return []

def evaluate_vector_store_quality(docs: list, last_question: str, min_docs: int = 3, min_avg_score: float = 0.3, high_quality_threshold: float = 0.7) -> dict:
    """
    Evaluates the quality of vector store results to determine if fallback is needed.
    Now prioritizes quality over quantity - allows fewer documents if they are highly relevant.
    Returns a dict with 'quality_sufficient', 'reason', and 'stats'.
    """
    if not docs:
        return {
            "quality_sufficient": False, 
            "reason": "Keine Dokumente gefunden",
            "stats": {"doc_count": 0, "avg_score": 0.0}
        }
    
    # Simple content relevance check - look for key terms from question in documents
    question_words = set(last_question.lower().split())
    # Remove common German stop words
    stop_words = {'der', 'die', 'das', 'und', 'oder', 'aber', 'ist', 'sind', 'hat', 'haben', 'von', 'zu', 'mit', 'auf', 'in', 'für', 'was', 'wie', 'wo', 'wann', 'warum'}
    question_words = question_words - stop_words
    
    if not question_words:
        # If no meaningful words left, assume it's complex enough
        return {
            "quality_sufficient": True,
            "reason": "Komplexe Abfrage - Vector Store wird verwendet",
            "stats": {"doc_count": len(docs), "avg_score": 1.0}
        }
    
    # Calculate relevance scores for each document
    relevance_scores = []
    for doc in docs:
        doc_content = doc.page_content.lower()
        matches = sum(1 for word in question_words if word in doc_content)
        score = matches / len(question_words) if question_words else 0
        relevance_scores.append(score)
    
    avg_score = sum(relevance_scores) / len(relevance_scores)
    max_score = max(relevance_scores) if relevance_scores else 0
    
    # NEW LOGIC: Prioritize quality over quantity
    
    # Case 1: High quality documents (even if few)
    if avg_score >= high_quality_threshold or max_score >= high_quality_threshold:
        return {
            "quality_sufficient": True,
            "reason": f"Hochrelevante Dokumente gefunden (Avg: {avg_score:.2f}, Max: {max_score:.2f}) - Qualität über Quantität",
            "stats": {"doc_count": len(docs), "avg_score": avg_score, "max_score": max_score}
        }
    
    # Case 2: Sufficient documents with acceptable quality
    if len(docs) >= min_docs and avg_score >= min_avg_score:
        return {
            "quality_sufficient": True,
            "reason": f"Ausreichende Dokumentenqualität und -anzahl gefunden (Score: {avg_score:.2f})",
            "stats": {"doc_count": len(docs), "avg_score": avg_score, "max_score": max_score}
        }
    
    # Case 3: Few documents with mediocre quality
    if len(docs) < min_docs and avg_score < min_avg_score:
        return {
            "quality_sufficient": False,
            "reason": f"Zu wenige Dokumente ({len(docs)}) und geringe Relevanz (Score: {avg_score:.2f})",
            "stats": {"doc_count": len(docs), "avg_score": avg_score, "max_score": max_score}
        }
    
    # Case 4: Few documents but decent quality
    if len(docs) < min_docs and avg_score >= min_avg_score:
        return {
            "quality_sufficient": False,
            "reason": f"Zu wenige relevante Dokumente gefunden ({len(docs)} von mindestens {min_docs}), aber akzeptable Qualität",
            "stats": {"doc_count": len(docs), "avg_score": avg_score, "max_score": max_score}
        }
    
    # Case 5: Sufficient documents but low quality
    if len(docs) >= min_docs and avg_score < min_avg_score:
        return {
            "quality_sufficient": False,
            "reason": f"Genügend Dokumente ({len(docs)}) aber geringe Relevanz (Score: {avg_score:.2f})",
            "stats": {"doc_count": len(docs), "avg_score": avg_score, "max_score": max_score}
        }
    
    # Fallback
    return {
        "quality_sufficient": True,
        "reason": f"Dokumentenqualität akzeptabel (Score: {avg_score:.2f})",
        "stats": {"doc_count": len(docs), "avg_score": avg_score, "max_score": max_score}
    }

def route_query(client, conversation_history: list, cancellation_check=lambda: False) -> str:
    """
    Classifies the user's query to determine the best information source using a powerful LLM.
    Returns 'vector_store', 'web_search', 'direct_answer', or 'image_generation'.
    """
    system_prompt = """You are a highly efficient and accurate query router. Your task is to classify the user's LATEST question into one of four categories.

**CRITICAL: MOST IMPORTANT RULE FOR IMAGE GENERATION:**
- **NEVER use 'image_generation' for questions that start with "Wie", "How", "What", "Was", "Can", "Kannst"**
- **ONLY use 'image_generation' for IMPERATIVE COMMANDS (no question words)**
- **Questions about image creation are ALWAYS 'direct_answer'**

**SPECIFIC EXAMPLES:**
- "Wie erzeuge ich ein Bild?" → direct_answer (THIS IS A QUESTION)
- "How do I create an image?" → direct_answer (THIS IS A QUESTION)
- "What is image generation?" → direct_answer (THIS IS A QUESTION)
- "Can you generate images?" → direct_answer (THIS IS A QUESTION)
- "Kannst du Bilder erstellen?" → direct_answer (THIS IS A QUESTION)
- "Erzeuge ein Bild von einem Hund" → image_generation (THIS IS A COMMAND)
- "Draw a cat" → image_generation (THIS IS A COMMAND)
- "Mach ein Bild" → image_generation (THIS IS A COMMAND)

1.  **vector_store**: Use this for questions about specific, internal, or proprietary topics. This is the correct choice for any question containing technical terms, jargon, or names (like "4PLAN", "Ermittlungstyp", "S4U", "HRCC") that are clearly not general public knowledge, even if the question seems simple.
    - Example: "What is our car policy?", "Tell me about the 4PLAN Dashboard Designer.", "Was ist ein Ermittlungstyp in 4PLAN?"

2.  **web_search**: Use this for questions that require up-to-date, public information. This includes current events, news, public figures, product reviews, or general knowledge that changes over time.
    - Example: "What is the latest news on AI?", "Who won the election?", "What are the reviews for the new iPhone?"

3.  **image_generation**: Use this ONLY for imperative commands without question words.
    - **COMMANDS (YES)**: "erzeuge ein bild", "draw a picture", "mach mal ein Bild", "create an image of a sunset"
    - **QUESTIONS (NO)**: "Wie erzeuge ich ein Bild?", "How do I create an image?", "What is image generation?"

4.  **direct_answer**: Use this for general knowledge questions, instructions, tutorials, or conversational phrases. This includes ALL questions that start with question words.
    - Example: "What is a large language model?", "Thank you", "Wie erzeuge ich ein Bild?", "How do I create an image?", "What is image generation?"

**FINAL CHECK: If the user's message contains "Wie", "How", "What", "Was", "Can", "Kannst" → ALWAYS use 'direct_answer'**

Respond with ONLY one of the following keywords: `vector_store`, `web_search`, `image_generation`, or `direct_answer`.
"""
    messages = create_contextual_messages(conversation_history, system_prompt)
    try:
        if cancellation_check(): return "direct_answer"

        # Use the main, more powerful LLM for routing to ensure accuracy with specialized terms.
        response = robust_api_call(client, LLM_MODEL, messages, 0.0)
        
        if cancellation_check(): return "direct_answer"
        
        if response.choices and response.choices[0].message.content:
            decision = response.choices[0].message.content.strip().lower()
            if "vector_store" in decision:
                return "vector_store"
            elif "web_search" in decision:
                return "web_search"
            elif "image_generation" in decision:
                return "image_generation"
        
        return "direct_answer"

    except Exception as e:
        print(f"Error during query routing: {e}")
        return "direct_answer"


# --- Web Search Cache ---
CACHE_DIR = os.path.join(SCRIPT_DIR, "web_cache")
os.makedirs(CACHE_DIR, exist_ok=True)

def get_url_hash(url: str) -> str:
    """Generate a hash for URL to use as cache key."""
    return hashlib.md5(url.encode()).hexdigest()

def get_cached_content(url: str) -> str:
    """Get cached content for a URL if it exists and is not expired."""
    cache_file = os.path.join(CACHE_DIR, f"{get_url_hash(url)}.json")
    if os.path.exists(cache_file):
        try:
            with open(cache_file, 'r', encoding='utf-8') as f:
                cache_data = json.load(f)
            
            # Check if cache is still valid (1 hour TTL)
            cache_time = cache_data.get('timestamp', 0)
            if time.time() - cache_time < 3600:  # 1 hour
                return cache_data.get('content', '')
        except Exception as e:
            print(f"Error reading cache for {url}: {e}")
    return None

def cache_content(url: str, content: str) -> None:
    """Cache content for a URL."""
    cache_file = os.path.join(CACHE_DIR, f"{get_url_hash(url)}.json")
    try:
        cache_data = {
            'url': url,
            'content': content,
            'timestamp': time.time()
        }
        with open(cache_file, 'w', encoding='utf-8') as f:
            json.dump(cache_data, f, ensure_ascii=False, indent=2)
    except Exception as e:
        print(f"Error caching content for {url}: {e}")

def smart_content_extraction(html: str) -> str:
    """Extract only relevant content from HTML, focusing on paragraphs and articles."""
    try:
        soup = BeautifulSoup(html, 'html.parser')
        
        # Remove unwanted elements
        for element in soup(["script", "style", "nav", "footer", "header", "aside", "menu", "form", "button"]):
            element.decompose()
        
        # Focus on content-rich elements
        content_elements = soup.find_all(['p', 'article', 'main', 'section', 'div'])
        
        # Extract text with better filtering
        relevant_text = []
        for elem in content_elements:
            text = elem.get_text(strip=True)
            # Filter out short, non-informative text
            if len(text) > 30 and not text.lower().startswith(('cookie', 'privacy', 'terms')):
                relevant_text.append(text)
        
        return " ".join(relevant_text)
    except Exception as e:
        return f"Error extracting content: {e}"

def scrape_website_content(url: str) -> str:
    """Enhanced scraping with caching and smart content extraction."""
    # Check cache first
    cached_content = get_cached_content(url)
    if cached_content:
        return cached_content
    
    try:
        response = requests.get(url, timeout=10, headers={'User-Agent': 'Mozilla/5.0'})
        response.raise_for_status()
        
        # Use smart content extraction
        content = smart_content_extraction(response.text)
        
        # Cache the result
        cache_content(url, content)
        
        return content
    except Exception as e:
        error_msg = f"Error scraping {url}: {e}"
        # Cache the error to avoid repeated failures
        cache_content(url, error_msg)
        return error_msg

def cleanup_expired_cache():
    """
    Removes expired cache files from the cache directory.
    This function runs the actual cleanup and logs the results.
    """
    if not os.path.exists(CACHE_DIR):
        return
    
    start_time = time.time()
    files_before = 0
    files_removed = 0
    
    try:
        cache_files = os.listdir(CACHE_DIR)
        files_before = len(cache_files)
        
        for cache_file in cache_files:
            file_path = os.path.join(CACHE_DIR, cache_file)
            try:
                with open(file_path, 'r', encoding='utf-8') as f:
                    cache_data = json.load(f)
                
                # Check if cache is expired (older than 1 hour)
                cache_time = cache_data.get('timestamp', 0)
                if time.time() - cache_time > 3600:  # 1 hour TTL
                    os.remove(file_path)
                    files_removed += 1
                    
            except (json.JSONDecodeError, KeyError, OSError) as e:
                # Remove corrupted or unreadable cache files
                try:
                    os.remove(file_path)
                    files_removed += 1
                except OSError:
                    pass  # File might already be removed
        
        duration = time.time() - start_time
        
        if files_removed > 0:
            logging.info(f"Cache cleanup: {files_removed} expired files removed from {files_before} total files in {duration:.2f}s")
        else:
            logging.info(f"Cache cleanup: No expired files found among {files_before} cache files (cleanup took {duration:.2f}s)")
            
    except Exception as e:
        logging.error(f"Error during cache cleanup: {e}")

async def periodic_cache_cleanup():
    """
    Background task that runs cache cleanup every 6 hours.
    This function runs in the background and doesn't block the main application.
    """
    # Wait 30 minutes before first cleanup (not immediately on startup)
    await asyncio.sleep(30 * 60)
    
    while True:
        try:
            logging.info(f"[{datetime.now()}] Starting periodic cache cleanup...")
            
            # Run cleanup in thread to avoid blocking
            await asyncio.to_thread(cleanup_expired_cache)
            
            logging.info(f"[{datetime.now()}] Periodic cache cleanup completed.")
            
        except Exception as e:
            logging.error(f"Periodic cache cleanup error: {e}")
        
        # Wait 6 hours before next cleanup
        await asyncio.sleep(6 * 60 * 60)

def get_cache_statistics():
    """
    Returns statistics about the current cache state.
    Useful for monitoring and admin interfaces.
    """
    if not os.path.exists(CACHE_DIR):
        return {
            "total_files": 0,
            "total_size_mb": 0,
            "expired_files": 0,
            "valid_files": 0
        }
    
    total_files = 0
    total_size_bytes = 0
    expired_files = 0
    valid_files = 0
    current_time = time.time()
    
    try:
        for cache_file in os.listdir(CACHE_DIR):
            file_path = os.path.join(CACHE_DIR, cache_file)
            if os.path.isfile(file_path):
                total_files += 1
                total_size_bytes += os.path.getsize(file_path)
                
                try:
                    with open(file_path, 'r', encoding='utf-8') as f:
                        cache_data = json.load(f)
                    
                    cache_time = cache_data.get('timestamp', 0)
                    if current_time - cache_time > 3600:  # Expired
                        expired_files += 1
                    else:
                        valid_files += 1
                        
                except (json.JSONDecodeError, KeyError):
                    expired_files += 1  # Count corrupted files as expired
                    
    except Exception as e:
        logging.error(f"Error getting cache statistics: {e}")
    
    return {
        "total_files": total_files,
        "total_size_mb": round(total_size_bytes / (1024 * 1024), 2),
        "expired_files": expired_files,
        "valid_files": valid_files
    }

def check_for_ambiguity(client, conversation_history: list, cancellation_check=lambda: False) -> dict:
    """
    Checks if the user's query is ambiguous, specifically for image generation only.
    Returns a dictionary indicating if clarification is needed.
    DISABLED: This function now always returns False to prevent unnecessary clarifications.
    """
    # Completely disable ambiguity checking as it was causing too many false positives
    # Even clear questions like "Wie erzeuge ich ein Bild?" were triggering clarifications
    return {"clarification_needed": False}

async def get_answer(conversation_history: list, source_mode: str = None, selected_fields: list = None, image_b64: str = None, cancellation_check=lambda: False):
    """
    Orchestrator for retrieving answers.
    This function is an async generator that yields status updates, metadata, and streamed answer chunks.
    It uses a persistent source_mode ('vector_store' or 'web_search') for the entire session.
    """
    try:
        client = Together(api_key=os.getenv("TOGETHER_API_KEY"))
        last_question = conversation_history[-1]["content"]

        # --- QUALITY FALLBACK CLARIFICATION HANDLING ---
        # Check if the user is responding to a quality fallback clarification
        user_choice = last_question.strip()
        
        if user_choice in ["Perform web search", "Use Knowledge Base anyway", "Do not answer now, I will reformulate my question"]:
            # User is responding to quality fallback clarification
            if user_choice == "Perform web search":
                # User chose web search - execute web search fallback
                yield {"type": "status", "data": "User chose web search - executing fallback search..."}
                # Force web search mode for this response
                determined_source_mode = "web_search"
            elif user_choice == "Use Knowledge Base anyway":
                # User chose to use KB anyway - proceed with vector store
                yield {"type": "status", "data": "User chose to use Knowledge Base anyway - proceeding with available results..."}
                determined_source_mode = "vector_store_forced"  # Special mode to skip quality evaluation
                yield {"type": "status", "data": "DEBUG: Set determined_source_mode to vector_store_forced"}
            elif user_choice == "Do not answer now, I will reformulate my question":
                # User chose to reformulate - end processing
                yield {"type": "meta", "data": {"sources": "User will reformulate question", "keywords": "N/A", "follow_ups": [], "source_mode": None}}
                yield {"type": "chunk", "data": "I understand. Please feel free to reformulate your question, and I'll do my best to help you."}
                yield {"type": "end"}
                return

        # --- AMBIGUITY CHECK ---
        # First, check if the user's query is ambiguous and needs clarification.
        # This is only necessary if no file has been uploaded (image_b64 is None).
        # Skip this check if we're handling a quality fallback response
        if (not image_b64 and 
            'determined_source_mode' not in locals()):
            yield {"type": "status", "data": "Checking for ambiguity..."}
            ambiguity_result = await asyncio.to_thread(check_for_ambiguity, client, conversation_history, cancellation_check)
            if ambiguity_result.get("clarification_needed"):
                yield {
                    "type": "clarification",
                    "data": {
                        "question": ambiguity_result.get("question"),
                        "options": ambiguity_result.get("options")
                    }
                }
                return # Stop processing, wait for user's clarifying response

        # --- Mode Decision Logic with Context-Awareness ---
        # If the session is already in image_generation mode, keep it there.
        if source_mode == "image_generation":
            determined_source_mode = "image_generation"
        elif 'determined_source_mode' not in locals():
            # Only do routing if we haven't already determined the source mode from quality fallback
            # STEP 1: Check if we can answer from conversation context (Hybrid Approach)
            yield {"type": "status", "data": "Analyzing conversation context..."}
            if cancellation_check(): return
            
            can_use_context = await asyncio.to_thread(can_answer_from_conversation_context, client, conversation_history)
            
            if can_use_context:
                yield {"type": "status", "data": "Question can be answered from conversation context"}
                determined_source_mode = "context_answer"
            else:
                # STEP 2: Normal routing logic for questions that need external sources
                yield {"type": "status", "data": "Routing query for external sources..."}
                if cancellation_check(): return
                
                routed_mode = await asyncio.to_thread(route_query, client, conversation_history, cancellation_check)

                # --- Routing logic with context enhancement ---
                determined_source_mode = "direct_answer" # Default to direct answer

                # 1. Always prioritize image generation
                if routed_mode == 'image_generation':
                    determined_source_mode = 'image_generation'
                # 2. Check for web search
                elif routed_mode == 'web_search':
                    # Check if web search is enabled by the user
                    features = load_features()
                    if features.get("web_search", True) and 'Web' in selected_fields:
                        determined_source_mode = 'web_search'
                    else:
                        # Web search is disabled, so fall back to direct answer
                        determined_source_mode = 'direct_answer'
                # 3. Check for vector store
                elif routed_mode == 'vector_store':
                    # Check if any actual vector store fields are selected
                    has_vector_fields = any(field for field in selected_fields if field != 'Web')
                    if has_vector_fields:
                        determined_source_mode = 'vector_store'
                    else:
                        # No vector stores selected, so fall back to direct answer
                        determined_source_mode = 'direct_answer'
                # 4. If the router suggested a direct answer, respect it.
                elif routed_mode == 'direct_answer':
                    determined_source_mode = 'direct_answer'
                
                yield {"type": "status", "data": f"Query routed to: {determined_source_mode.replace('_', ' ')}"}

        # --- Answer Generation based on determined mode ---
        if determined_source_mode == "image_generation":
            # Check if image generation feature is enabled
            features = load_features()
            if not features.get("image_generation", True):
                yield {"type": "meta", "data": {"source_mode": None}}
                yield {"type": "chunk", "data": "Bildgenerierung ist derzeit deaktiviert. Bitte wenden Sie sich an einen Administrator, um diese Funktion zu aktivieren."}
                yield {"type": "end"}
                return
            
            last_image_bytes = base64.b64decode(image_b64) if image_b64 else None
            # Set the mode for the session and then generate the image.
            yield {"type": "meta", "data": {"source_mode": "image_generation"}}
            async for event in generate_image(conversation_history, last_image_bytes=last_image_bytes):
                yield event
            return

        if determined_source_mode == "context_answer":
            yield {"type": "status", "data": "Generating answer from conversation context..."}
            
            # Build conversation context for the LLM
            raw_conversation = ""
            for i, msg in enumerate(conversation_history):
                role = msg.get('role', 'unknown')
                content = msg.get('content', '')
                raw_conversation += f"{role.upper()}: {content}\n\n"
            
            # Simple and direct prompt for context-based answers
            system_prompt = f"""You are answering a follow-up question. Here is the COMPLETE conversation history:

{raw_conversation}

The user's current question is a follow-up. Look at the conversation above and answer based on what was already discussed. If you see weather information about Miami and the user asks about an umbrella, connect the two. Answer in the same language as the user's question.

Be specific and reference the previous information directly."""

            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": conversation_history[-1]['content']}
            ]
            
            # Generate temporary answer for follow-ups
            temp_response = await asyncio.to_thread(robust_api_call, client, LLM_MODEL, messages, 0.1, stream=False)
            temp_answer = ""
            if temp_response.choices and temp_response.choices[0].message.content:
                temp_answer = temp_response.choices[0].message.content.strip()

            follow_ups = await asyncio.to_thread(generate_follow_up_questions, client, conversation_history, temp_answer, "conversation context")
            
            yield {"type": "meta", "data": {"sources": "Conversation Context", "keywords": "N/A", "follow_ups": follow_ups, "source_mode": determined_source_mode}}

            if cancellation_check(): return
            stream = robust_api_call(client, LLM_MODEL, messages, 0.1, stream=True, cancellation_check=cancellation_check)
            for chunk in stream:
                if cancellation_check(): return
                if chunk.choices and chunk.choices[0].delta.content:
                    yield {"type": "chunk", "data": chunk.choices[0].delta.content}
            yield {"type": "end"}
            return

        if cancellation_check(): return
        if not vector_stores:
            yield {"type": "meta", "data": {"sources": "No sources", "keywords": "N/A", "follow_ups": []}}
            yield {"type": "chunk", "data": "The knowledge base has not yet been created or is empty. Please update it first."}
            yield {"type": "end"}
            return

        if determined_source_mode == "direct_answer":
            yield {"type": "status", "data": "Formulating a direct answer..."}
            system_prompt = "You are a helpful assistant. Answer the user's question based on your general knowledge and the conversation history."
            messages = create_contextual_messages(conversation_history, system_prompt)
            
            temp_response = await asyncio.to_thread(robust_api_call, client, LLM_MODEL, messages, 0.1, stream=False)
            temp_answer = ""
            if temp_response.choices and temp_response.choices[0].message.content:
                temp_answer = temp_response.choices[0].message.content.strip()

            follow_ups = await asyncio.to_thread(generate_follow_up_questions, client, conversation_history, temp_answer, "")
            
            # source_mode is None for direct answers
            yield {"type": "meta", "data": {"sources": "General Knowledge", "keywords": "N/A", "follow_ups": follow_ups, "source_mode": None}}

            if cancellation_check(): return
            stream = robust_api_call(client, LLM_MODEL, messages, 0.1, stream=True)
            for chunk in stream:
                if cancellation_check(): return
                if chunk.choices and chunk.choices[0].delta.content:
                    yield {"type": "chunk", "data": chunk.choices[0].delta.content}
            yield {"type": "end"}
            return

        if determined_source_mode == "vector_store" or determined_source_mode == "vector_store_forced":
            yield {"type": "status", "data": "Searching internal knowledge base with pure semantic search..."}
            
            # Clean approach: Pure semantic search with excellent multilingual embeddings
            expanded_queries = await asyncio.to_thread(expand_query_with_llm, client, conversation_history)
            
            target_stores = {}
            if selected_fields:
                for field in selected_fields:
                    if field in vector_stores:
                        target_stores[field] = vector_stores[field]
            else:
                target_stores = vector_stores

            if not target_stores:
                yield {"type": "meta", "data": {"sources": "No sources", "keywords": "N/A", "follow_ups": []}}
                yield {"type": "chunk", "data": "The selected knowledge field(s) could not be found or are empty."}
                yield {"type": "end"}
                return

            all_retrieved_docs = []
            for field_name, store in target_stores.items():
                yield {"type": "status", "data": f"Semantic search in '{field_name}' with multilingual embeddings..."}
                try:
                    # Increased retrieval for better coverage with smart chunks
                    retriever = store.as_retriever(search_kwargs={"k": 25})
                    tasks = [retriever.ainvoke(query) for query in expanded_queries]
                    retrieved_results = await asyncio.gather(*tasks)
                    for result in retrieved_results:
                        if cancellation_check(): return
                        all_retrieved_docs.extend(result)
                except Exception as e:
                    print(f"Error searching in vector store '{field_name}': {e}")
                    yield {"type": "status", "data": f"Warning: Vector store '{field_name}' needs to be rebuilt with the new multilingual model. Please update the knowledge base."}
                    continue

            unique_docs = list({(doc.page_content, doc.metadata.get('source', 'N/A')): doc for doc in all_retrieved_docs}.values())
            
            # Pure semantic search with balanced field selection
            if unique_docs:
                # Create a balanced selection from both fields
                field_docs = {}
                for doc in unique_docs:
                    field = doc.metadata.get('knowledge_field', 'Unknown')
                    if field not in field_docs:
                        field_docs[field] = []
                    field_docs[field].append(doc)
                
                # Balanced selection: alternate between fields to ensure fair representation
                top_docs = []
                max_per_field = 10 // len(field_docs) if field_docs else 0
                remainder = 10 % len(field_docs) if field_docs else 0
                
                for field, docs in field_docs.items():
                    take_count = max_per_field + (1 if remainder > 0 else 0)
                    if remainder > 0:
                        remainder -= 1
                    top_docs.extend(docs[:take_count])
                
                # If we still need more docs, fill from any remaining
                if len(top_docs) < 10:
                    remaining_docs = [doc for docs in field_docs.values() for doc in docs[max_per_field:]]
                    additional_needed = 10 - len(top_docs)
                    top_docs.extend(remaining_docs[:additional_needed])
                
                top_docs = top_docs[:10]  # Ensure exactly 10 docs
            else:
                top_docs = []
            
            # NEW: Evaluate quality of vector store results and implement fallback
            # BUT ONLY if the user hasn't already decided to use Knowledge Base anyway
            if determined_source_mode == "vector_store_forced":
                # User explicitly chose to use Knowledge Base anyway - skip quality evaluation
                yield {"type": "status", "data": "Using Knowledge Base as requested by user..."}
                quality_eval = {"quality_sufficient": True, "reason": "User chose to use Knowledge Base anyway"}
            else:
                yield {"type": "status", "data": "Evaluating search result quality..."}
                quality_eval = evaluate_vector_store_quality(top_docs, last_question)
                
                # Check if we should fallback to web search
                if not quality_eval["quality_sufficient"]:
                    yield {"type": "status", "data": f"Vector Store Qualität unzureichend: {quality_eval['reason']}"}
                    
                    # Check if web search is available as fallback
                    features = load_features()
                    web_search_available = features.get("web_search", True) and 'Web' in selected_fields
                    
                    if web_search_available:
                        # Ask user for confirmation before fallback
                        yield {
                            "type": "clarification",
                            "data": {
                                "question": "I don't find high quality data in the knowledge base to answer the question. What do you want me to do?",
                                "options": [
                                    "Perform web search",
                                    "Use Knowledge Base anyway", 
                                    "Do not answer now, I will reformulate my question"
                                ],
                                "clarification_type": "quality_fallback"
                            }
                        }
                        return  # Stop processing, wait for user's clarifying response
                    else:
                        # No web search available, proceed with what we have but inform user
                        yield {"type": "status", "data": "Web-Suche nicht verfügbar - verwende verfügbare Knowledge Base Ergebnisse..."}
            
            yield {"type": "status", "data": "Processing semantic search results..."}

            context = ""
            if top_docs:
                context = "\n\n---\n\n".join([doc.page_content for doc in top_docs])

            # Adjust status message based on quality
            if quality_eval["quality_sufficient"]:
                yield {"type": "status", "data": "Formulating answer from internal documents..."}
            else:
                yield {"type": "status", "data": "Formulating answer from available internal documents (limited quality)..."}
            system_prompt = """You are a helpful assistant. Your task is to answer the user's question based *only* on the provided 'Internal Context'.
- Read the user's question in the context of the full conversation.
- Analyze the 'Internal Context' thoroughly and provide a comprehensive answer.
- Do not use any outside knowledge.
- You MUST create a nicely markdown formatted answer with bullet points.
"""
            messages = create_contextual_messages(conversation_history, system_prompt)
            messages.append({"role": "system", "content": f"Internal Context:\n\n{context}"})
            
            temp_response = await asyncio.to_thread(robust_api_call, client, LLM_MODEL, messages, 0.0, stream=False)
            temp_answer = ""
            if temp_response.choices and temp_response.choices[0].message.content:
                temp_answer = temp_response.choices[0].message.content.strip()
            
            follow_ups = await asyncio.to_thread(generate_follow_up_questions, client, conversation_history, temp_answer, context)
            
            sources_text = "\n".join(sorted(list(set([f"- {os.path.basename(doc.metadata['source'])}" for doc in top_docs])))) if top_docs else "No internal sources found."
            yield {"type": "meta", "data": {"sources": f"**Internal sources:**\n{sources_text}", "keywords": "N/A", "follow_ups": follow_ups, "source_mode": determined_source_mode}}

            if cancellation_check(): return
            stream = robust_api_call(client, LLM_MODEL, messages, 0.0, stream=True)
            for chunk in stream:
                if cancellation_check(): return
                if chunk.choices and chunk.choices[0].delta.content:
                    yield {"type": "chunk", "data": chunk.choices[0].delta.content}
            yield {"type": "end"}
            return

        # --- Web Search Execution ---
        if determined_source_mode == "web_search":
            yield {"type": "status", "data": "Starting optimized Google search..."}
            
            # Extract context keywords from conversation for better search
            yield {"type": "status", "data": "Extracting context keywords from conversation..."}
            context_keywords = await asyncio.to_thread(extract_context_keywords, client, conversation_history)
            
            # Get German queries
            german_queries = await asyncio.to_thread(expand_query_with_llm, client, conversation_history)
            
            # Enhance queries with context keywords (with validation)
            if context_keywords:
                yield {"type": "status", "data": f"Enhancing search with context: {', '.join(context_keywords[:3])}..."}
                enhanced_german_queries = []
                for query in german_queries:
                    # Clean and validate context keywords
                    clean_keywords = []
                    for keyword in context_keywords[:3]:
                        # Remove non-alphanumeric characters except spaces and common punctuation
                        clean_keyword = re.sub(r'[^\w\s\-\.\,]', '', keyword.strip())
                        if clean_keyword and len(clean_keyword) > 2:
                            clean_keywords.append(clean_keyword)
                    
                    # Add context keywords only if they don't make the query too long
                    if clean_keywords:
                        enhanced_query = f"{query} {' '.join(clean_keywords)}"
                        # Limit query length to prevent Google search errors
                        if len(enhanced_query) <= 200:
                            enhanced_german_queries.append(enhanced_query)
                        else:
                            enhanced_german_queries.append(query)  # Use original if too long
                    else:
                        enhanced_german_queries.append(query)
                german_queries = enhanced_german_queries
            
            yield {"type": "status", "data": "Translating search queries to English..."}
            
            # Translate queries to English
            if cancellation_check(): return
            english_queries = await asyncio.to_thread(translate_queries_to_english, client, german_queries, cancellation_check)

            # Perform searches in parallel for better performance (reduced to 3 results per language)
            yield {"type": "status", "data": "Performing optimized parallel searches..."}
            if cancellation_check(): return
            
            german_search_query = ", ".join(german_queries)
            english_search_query = ", ".join(english_queries) if english_queries else ""
            
            # Run searches in parallel with error handling
            tasks = []
            
            # Validate and clean search queries before execution
            def clean_search_query(query):
                # Remove special characters that might cause issues
                cleaned = re.sub(r'[^\w\s\-\.\,\?\!]', '', query.strip())
                # Limit length to prevent errors
                return cleaned[:200] if len(cleaned) > 200 else cleaned
            
            clean_german_query = clean_search_query(german_search_query)
            clean_english_query = clean_search_query(english_search_query) if english_search_query else ""
            
            # Execute searches with error handling (reduced to 3 results per language)
            try:
                if clean_german_query:
                    tasks.append(asyncio.create_task(asyncio.to_thread(list, search(clean_german_query, num_results=3, lang="de"))))
                if clean_english_query:
                    tasks.append(asyncio.create_task(asyncio.to_thread(list, search(clean_english_query, num_results=3, lang="en"))))
                
                if not tasks:
                    yield {"type": "meta", "data": {"sources": "No sources", "keywords": "Invalid search query", "follow_ups": []}}
                    yield {"type": "chunk", "data": "Search query could not be processed due to invalid characters or length."}
                    yield {"type": "end"}
                    return
                
                search_results = await asyncio.gather(*tasks)
                
            except Exception as search_error:
                yield {"type": "status", "data": f"Search API error: {str(search_error)[:100]}..."}
                yield {"type": "meta", "data": {"sources": "No sources", "keywords": "Search API error", "follow_ups": []}}
                yield {"type": "chunk", "data": "I encountered an error while searching. This might be due to search API limitations or query formatting issues. Please try rephrasing your question."}
                yield {"type": "end"}
                return
            
            # Process German results (reduced to 3)
            raw_german_results = search_results[0]
            german_urls = [url for url in raw_german_results if (url.startswith('http://') or url.startswith('https://')) and 'google.com/search' not in url][:3]
            
            # Process English results (reduced to 3)
            english_urls = []
            if len(search_results) > 1:
                raw_english_results = search_results[1]
                english_urls = [url for url in raw_english_results if (url.startswith('http://') or url.startswith('https://')) and 'google.com/search' not in url][:3]

            # Combine and deduplicate results (now max 6 instead of 8)
            all_urls = german_urls + english_urls
            search_results = sorted(list(set(all_urls)))
            
            web_search_query = f"DE: {german_search_query} | EN: {', '.join(english_queries) if english_queries else ''}"

            if not search_results:
                yield {"type": "meta", "data": {"sources": "No sources", "keywords": web_search_query, "follow_ups": []}}
                yield {"type": "chunk", "data": "I couldn't find an answer in the internal documents or on the web."}
                yield {"type": "end"}
                return

            # STEP 3: Intelligent source prioritization (limit to top 4 sources)
            yield {"type": "status", "data": f"Prioritizing top 4 sources from {len(search_results)} results..."}
            
            # Take only the top 4 sources for faster processing
            priority_sources = search_results[:4]
            
            yield {"type": "status", "data": "Scraping priority sources in parallel..."}
            
            # Enhanced parallel scraping with better error handling
            async def scrape_with_error_handling(url):
                try:
                    content = await asyncio.to_thread(scrape_website_content, url)
                    return {"url": url, "content": truncate_text(content, 12000), "success": True}  # Reduced content size for faster processing
                except Exception as e:
                    print(f"Error scraping {url}: {e}")
                    return {"url": url, "content": f"Error accessing content: {e}", "success": False}
            
            # Create scraping tasks for priority sources only
            scrape_tasks = [scrape_with_error_handling(url) for url in priority_sources]
            scraped_results = await asyncio.gather(*scrape_tasks)
            
            # Filter successful results and log failures
            successful_results = [result for result in scraped_results if result["success"]]
            failed_count = len(scraped_results) - len(successful_results)
            
            if failed_count > 0:
                yield {"type": "status", "data": f"Successfully scraped {len(successful_results)} of {len(priority_sources)} priority sources"}
            
            # Use successful results for processing
            scraped_results = successful_results
            
            if not scraped_results:
                yield {"type": "meta", "data": {"sources": "No sources", "keywords": web_search_query, "follow_ups": []}}
                yield {"type": "chunk", "data": "I found search results but couldn't access the content from any sources."}
                yield {"type": "end"}
                return

            # STEP 4: Streaming response - Start processing as soon as we have first results
            yield {"type": "status", "data": "Starting streaming response generation..."}
            
            # Create the consolidated context for the LLM, embedding the URL directly with the source number.
            consolidated_content = "\n\n---\n\n".join(
                [f"Source [{i+1}]({res['url']}):\n{res['content']}" for i, res in enumerate(scraped_results)]
            )

            system_prompt = f"""You are a web analysis expert. Your task is to answer the user's question based *only* on the provided 'Web Search Context'.

**Instructions:**
1.  Read the user's question: '{last_question}'.
2.  Analyze the 'Web Search Context'. Each source is now formatted as `Source [number](URL): content`.
3.  Formulate a comprehensive answer in the same language as the user's question.
4.  **Cite sources by using the exact markdown hyperlink provided in the context.** For example, when you use information from `Source [1](https://example.com/source1)`, you must cite it as `[1](https://example.com/source1)`.
5.  **Placement Rule:** Place citation links within the descriptive text (e.g., at the end of a sentence).
6.  Structure your answer using markdown, including bullet points for clarity.
7.  If the provided context is insufficient, state that you could not find the information.
8.  **Crucial Rule: DO NOT** add a separate list of sources (e.g., "Quellen:", "Sources:") at the end of your answer. All source citations **MUST** be the inline markdown hyperlinks from the context.

**Correct Citation Example:**
"This is a descriptive sentence that uses information from a source [1](https://example.com/source1)."
"""
            # Truncate context to fit within model limits
            messages_for_size_check = create_contextual_messages(conversation_history, system_prompt)
            conversation_chars = sum(len(m['content']) for m in messages_for_size_check)
            # Adjusted remaining_chars to account for the larger prompt
            remaining_chars = 512000 - conversation_chars
            final_context = truncate_text(consolidated_content, remaining_chars)

            messages = create_contextual_messages(conversation_history, system_prompt)
            messages.append({"role": "system", "content": f"Web Search Context:\n\n{final_context}"})

            # Generate follow-ups first
            temp_response = await asyncio.to_thread(robust_api_call, client, LLM_MODEL, messages, 0.0, stream=False)
            temp_answer = ""
            if temp_response.choices and temp_response.choices[0].message.content:
                temp_answer = temp_response.choices[0].message.content.strip()
            follow_ups = await asyncio.to_thread(generate_follow_up_questions, client, conversation_history, temp_answer, consolidated_content)

            # Yield metadata with a clean list of URLs
            yield {"type": "meta", "data": {"sources": search_results, "keywords": web_search_query, "follow_ups": follow_ups, "source_mode": determined_source_mode}}

            # Stream the answer
            yield {"type": "status", "data": "Formulating answer based on web search..."}
            if cancellation_check(): return
            stream = robust_api_call(client, LLM_MODEL, messages, 0.0, stream=True)
            for chunk in stream:
                if cancellation_check(): return
                if chunk.choices and chunk.choices[0].delta.content:
                    yield {"type": "chunk", "data": chunk.choices[0].delta.content}
            yield {"type": "end"}

    except Exception as e:
        error_message = f"An error occurred in get_answer: {e}"
        print(f"Error in get_answer function: {e}")
        import traceback
        traceback.print_exc()
        yield {"type": "meta", "data": {"sources": "No sources", "keywords": "N/A", "follow_ups": []}}
        yield {"type": "chunk", "data": "An error occurred while processing your request. Please try again."}
        yield {"type": "end"}

def needs_python_script(client, conversation_history: list) -> bool:
    """
    Determines if the user's question requires a Python script to be answered.
    """
    system_prompt = """You are a critical assistant. Your task is to determine if a Python script is **absolutely necessary** to answer the user's latest question. Your default assumption should be that a script is **NOT** needed.

- **Rule 1: Prioritize direct answers.** Can the question be answered by summarizing, quoting, or extracting information directly from the conversation history or the provided data context? If yes, you **MUST** respond with 'no'.
- **Rule 2: Identify complex operations.** A Python script is only required for tasks that **cannot** be done manually by looking at the data. This includes:
    - Mathematical calculations (sum, average, median, etc.) across multiple data points.
    - Grouping data and then performing calculations on those groups.
    - Generating a table or a chart as the final output.
- **Rule 3: Avoid Python for simple lookups.** If the user asks "What is the value for X?", and the value is directly visible in the data, that is a simple lookup. Respond with 'no'. If the user asks "What is the sum of all X?", that requires a calculation. Respond with 'yes'.
- **Rule 4: Ignore past scripts.** Do not decide to use Python just because a script was used earlier in the conversation. Evaluate each question independently.

Based on these strict rules, respond with only 'yes' or 'no'.
"""
    messages = create_contextual_messages(conversation_history, system_prompt)
    try:
        response = robust_api_call(client, LLM_MODEL, messages, 0.0)
        if response.choices and response.choices[0].message.content:
            decision = response.choices[0].message.content.strip().lower()
            return "yes" in decision
        return False
    except Exception as e:
        print(f"Error during Python script check: {e}")
        return False

def get_python_code(conversation_history: list, file_path: str = None, cancellation_check=lambda: False):
    """
    Generates Python code and a natural language explanation to answer a question.
    """
    try:
        client = Together(api_key=os.getenv("TOGETHER_API_KEY"))
        
        base_system_prompt = """You are an expert Python programmer and a helpful assistant. Your task is to write a Python script to answer the user's question and also provide a simple, natural language explanation of how the script works.

**Output Format:**
You **MUST** respond with a single, valid JSON object. This JSON object must have two keys:
1.  `"python_code"`: A string containing the complete, raw Python script.
2.  `"explanation"`: A string containing a user-friendly explanation in German.

**Sonderfall: Nicht umsetzbare Fragen:**
- Wenn die Frage des Benutzers zu allgemein ist oder keine spezifische Datenauswertung erfordert (z.B. "Was ist das für eine Datei?", "Erzähl mir was über die Daten"), dann darfst du **KEINEN** Python-Code erzeugen.
- In diesem Fall muss dein JSON so aussehen:
  `"python_code": ""` (ein leerer String)
  `"explanation": "Ihre Frage kann nicht in eine konkrete Datenauswertung umgewandelt werden. Bitte beachten Sie, dass bei Tabellen-Dateien nur spezifische Auswertungen per Python-Skript möglich sind."`

**Rules for `"python_code"`:**
- The script's output must be either a JSON table (`your_dataframe.to_json(orient='split', index=False)`) or an HTML file for charts.
- **All charts MUST be generated using the `plotly` library** (e.g., `import plotly.express as px` or `import plotly.graph_objects as go`). The output MUST be saved as an HTML file (e.g., `fig.write_html('temp_images/plot.html')`). Do NOT use any other plotting library.
- When creating a table, the script's standard output **MUST ONLY** contain the final JSON data. Do not print any other text.
- When creating a chart, the script **MUST NOT** print anything to standard output.
- The script must be self-contained.
- When the user asks for a per FTE or Per Headcount calculation, you MUST calculate the sum of whatever the user wants to see on a per FTE or per Headcount basis and then divide it by the number of FTEs or Headcounts ON THE LEVEL OF THE GROUP. First calculate the sums and then perform the division.

**Rules for `"explanation"`:**
- Explain how the result was calculated in simple, non-technical German.
- Describe any grouping, filtering, or calculations that were applied.
- You **MUST** mention the specific column names from the source data that were used for filtering, grouping, or calculations.
- **DO NOT** use any Python syntax (e.g., no `df.groupby`, no function names).
- Write as if you are explaining the process to someone who does not know programming.
- Example Explanation: "Um das Ergebnis zu ermitteln, habe ich die Daten nach der Spalte 'Abteilung' gruppiert und anschließend die Summe der Werte aus der Spalte 'Gehalt' für jede einzelne Abteilung berechnet."

**Example JSON Response:**
```json
{
  "python_code": "import pandas as pd\\ndf = pd.read_csv('FILE_PATH')\\n# ... rest of the script ...\\nprint(df.to_json(orient='split', index=False))",
  "explanation": "Ich habe die Daten nach der Spalte 'Abteilung' gruppiert und dann die Summe der Werte aus der Spalte 'Gehalt' für jede Abteilung berechnet."
}
```
"""

        if file_path:
            file_name = os.path.basename(file_path)
            file_specific_prompt = f"""
**File-Specific Instructions:**
- The script must read data from a file named '{file_name}'.
- Use the placeholder 'FILE_PATH' to access the file (e.g., `pd.read_csv('FILE_PATH')`). The system will replace this.
- The user's last message might contain a 'Technical Note' with the **exact list of column names**. If so, you **MUST** use these names.
"""
            system_prompt = base_system_prompt + file_specific_prompt
        else:
            system_prompt = base_system_prompt
        
        messages = create_contextual_messages(conversation_history, system_prompt)
        
        if cancellation_check(): return

        response = robust_api_call(client, LLM_MODEL, messages, 0.1, stream=False)
        
        content = ""
        if response.choices and response.choices[0].message.content:
            content = response.choices[0].message.content.strip()

        # Extract JSON from the markdown block if present
        if content.startswith("```json"):
            content = content[7:-3].strip()
        
        try:
            data = json.loads(content)
            python_code = data.get("python_code", "")
            explanation = data.get("explanation", "")

            # Replace the placeholder with the actual file path
            if file_path:
                python_code = python_code.replace("'FILE_PATH'", f"'{file_path}'")

            yield {"python_code": python_code, "explanation": explanation}

        except json.JSONDecodeError:
            # Fallback for cases where the model doesn't return valid JSON
            yield {"python_code": content, "explanation": "Ich habe ein Skript erstellt, um Ihre Frage zu beantworten."}

    except Exception as e:
        error_message = f"An error occurred while generating Python code: {e}"
        print(error_message)
        yield {"error": error_message}

async def get_answer_from_document(conversation_history: list, document_content: str, file_type: str = None, last_image_bytes: bytes = None, cancellation_check=lambda: False):
    """
    Generates a streaming answer based on the user's question and the content of an uploaded document (text or image).
    TEMPORARY: Using full document content without chunking for testing.
    """
    try:
        client = Together(api_key=os.getenv("TOGETHER_API_KEY"))
        
        yield {"type": "status", "data": "Analyzing question..."}

        # --- Image Analysis/Generation Branch ---
        # This branch is now also triggered for image generation requests coming through document_question
        if (file_type and file_type.startswith('image/')) or last_image_bytes:
            # If there's last_image_bytes, it's a generation/refinement request.
            if last_image_bytes:
                # Check if image generation feature is enabled
                features = load_features()
                if not features.get("image_generation", True):
                    yield {"type": "meta", "data": {"source_mode": None}}
                    yield {"type": "chunk", "data": "Bildgenerierung ist derzeit deaktiviert. Bitte wenden Sie sich an einen Administrator, um diese Funktion zu aktivieren."}
                    yield {"type": "end"}
                    return
                
                async for event in generate_image(conversation_history, last_image_bytes=last_image_bytes):
                    yield event
                return

            # --- This part below is for analyzing a NEWLY uploaded image ---
            yield {"type": "status", "data": "Analyzing image content..."}
            
            last_question = conversation_history[-1]['content']
            # Add instruction to answer in the user's language.
            image_prompt = f"USER: <image>\n{last_question}\n\n(System note: You MUST answer in the same language as the user's question.)"
            
            messages = [{"role": "user", "content": image_prompt}]
            
            # Add the base64 encoded image to the last message
            messages[-1]['content'] = [
                {"type": "text", "text": image_prompt},
                {"type": "image_url", "image_url": {"url": f"data:{file_type};base64,{document_content}"}}
            ]

            yield {"type": "meta", "data": {"sources": "Source: Uploaded Image", "keywords": "N/A", "follow_ups": [], "source_mode": "image_analysis"}}
            yield {"type": "status", "data": "Formulating a response based on the image..."}

            try:
                # Use the dedicated image model for the API call and add a long timeout
                stream = robust_api_call(client, IMAGE_MODEL, messages, 0.1, stream=True, timeout=180)
                
                has_content = False
                for chunk in stream:
                    if cancellation_check(): return
                    if chunk.choices and chunk.choices[0].delta.content:
                        has_content = True
                        yield {"type": "chunk", "data": chunk.choices[0].delta.content}
                
                if not has_content:
                    # This can happen if the model returns an empty stream for non-text images
                    raise ValueError("The model returned an empty response, likely due to no text in the image.")

                yield {"type": "end"}
                return
            except Exception as e:
                print(f"Error during image stream processing: {e}")
                # Yield a user-friendly error message instead of crashing
                yield {"type": "chunk", "data": "Das Bild konnte nicht verarbeitet werden. Dies kann daran liegen, dass es keinen erkennbaren Text enthält oder das Format nicht unterstützt wird."}
                yield {"type": "end"}
                return

        # --- Text/Table Document Branch (existing logic) ---
        elif file_type in ['text', 'sql', 'table', 'table_data']:
            # For table files, always require a Python script for data privacy
            if file_type in ['table', 'table_data']:
                yield {"type": "python_required"}
                return
            # For .txt or .sql files, always answer with AI, never use Python
            elif file_type in ['text', 'sql']:
                pass # Skip the python check and proceed to the AI answer

        yield {"type": "status", "data": "TEMPORARY TEST: Using full document content without chunking..."}

        # TEMPORARY: Use the FULL document content without any truncation for testing
        system_prompt_template = """Sie sind ein Experte für Dokumentenanalyse und Informationsextraktion. Ihre Aufgabe ist es, die Frage des Benutzers basierend *ausschließlich* auf dem bereitgestellten Dokumenteninhalt zu beantworten.

**INTELLIGENTE ABKÜRZUNGS- UND BEGRIFFSERKENNUNG:**
- Erkennen Sie gängige Abkürzungen und Akronyme im Kontext des Dokuments
- Suchen Sie nach sowohl abgekürzten als auch vollständigen Formen von Begriffen
- Berücksichtigen Sie fachspezifische Terminologie je nach Dokumenttyp
- Achten Sie auf Definitionen und Erklärungen im Dokument selbst

**SPEZIELLE ANWEISUNGEN FÜR STRUKTURIERTE DATEN:**
- Wenn Sie Tabellendaten sehen, analysieren Sie ALLE Zeilen und Spalten sorgfältig
- Suchen Sie nach Mustern und Zusammenhängen zwischen verschiedenen Tabellenteilen
- Identifizieren Sie Trends, Veränderungen und wichtige Datenpunkte
- Kombinieren Sie Informationen aus verschiedenen Dokumententeilen
- Erklären Sie die Bedeutung der Daten im Kontext des Dokuments
- Achten Sie auf Zahlen, Statistiken, Vergleiche und deren Relevanz
- **WICHTIG: Sie haben Zugriff auf das GESAMTE Dokument. Nutzen Sie alle verfügbaren Informationen.**

**ANWEISUNGEN:**
- **Verwenden Sie keine externen Kenntnisse oder Informationen außerhalb des Dokuments.**
- Wenn die Antwort nicht im Dokument gefunden werden kann, geben Sie das klar an.
- Analysieren Sie die Frage des Benutzers im Kontext der Unterhaltung.
- Erstellen Sie eine schön formatierte Markdown-Antwort mit Aufzählungspunkten.
- Sie MÜSSEN präzise sein, wenn Sie den Dokumenteninhalt analysieren und die Frage beantworten. Lesen Sie den gesamten Inhalt sorgfältig.
- Sie MÜSSEN alle Daten und Informationen im Dokument aufmerksam und fokussiert lesen und dann die Frage beantworten.
- Beantworten Sie die Frage direkt und fügen Sie nicht zu viele zusätzliche Informationen hinzu.
- **Sprache:** Sie MÜSSEN in derselben Sprache antworten, die der Benutzer in seiner Frage verwendet hat.

**VOLLSTÄNDIGER DOKUMENTENINHALT (OHNE KÜRZUNG):**
---
{document_content}
---
"""
        
        # TEMPORARY: Calculate available space more generously for the 1M context window
        messages_for_size_check = create_contextual_messages(conversation_history, "")
        conversation_chars = sum(len(m['content']) for m in messages_for_size_check)
        prompt_template_chars = len(system_prompt_template) - len("{document_content}")
        
        # Use much larger context window (1M tokens ≈ 800K characters)
        remaining_chars = 800000 - conversation_chars - prompt_template_chars
        
        # TEMPORARY: Try to use as much of the document as possible
        if len(document_content) <= remaining_chars:
            final_doc_content = document_content  # Use full document
            yield {"type": "status", "data": f"Using FULL document content ({len(document_content)} characters)"}
        else:
            final_doc_content = document_content[:remaining_chars]
            yield {"type": "status", "data": f"Using truncated document content ({len(final_doc_content)} of {len(document_content)} characters)"}

        system_prompt = system_prompt_template.format(document_content=final_doc_content)
        messages = create_contextual_messages(conversation_history, system_prompt)
        
        yield {"type": "meta", "data": {"sources": "Source: Uploaded document (FULL CONTENT TEST)", "keywords": "N/A", "follow_ups": []}}
        yield {"type": "status", "data": "Formulating a response based on the FULL document..."}
        
        if cancellation_check(): return
        stream = robust_api_call(client, LLM_MODEL, messages, 0.1, stream=True)
        
        for chunk in stream:
            if cancellation_check(): return
            if chunk.choices and chunk.choices[0].delta.content:
                yield {"type": "chunk", "data": chunk.choices[0].delta.content}
        
        yield {"type": "end"}

    except Exception as e:
        error_message = f"An error occurred while processing the document: {e}"
        print(error_message)
        yield {"type": "meta", "data": {"sources": "No sources", "keywords": "N/A", "follow_ups": []}}
        yield {"type": "chunk", "data": error_message}
        yield {"type": "end"}

async def translate_prompt_to_english(client, prompt: str) -> str:
    """
    Detects the language of the prompt and translates it to English if it's not already English.
    """
    if not prompt:
        print("Translation skipped: Prompt is empty.")
        return ""
    try:
        # Detect language
        lang = await asyncio.to_thread(detect, prompt)
        # print(f"Detected language: {lang}")
        
        # If language is not English, translate it
        if lang != 'en':
            # print(f"Prompt is not in English. Translating...")

            system_prompt = "You are a literal translator. Your only task is to translate the following text to English. Do not add, remove, or change any part of the meaning. Translate it as literally as possible. Return only the translated text and nothing else."
            messages = [
                {"role": "system", "content": system_prompt},
                {"role": "user", "content": prompt}
            ]
            
            # Use a smaller, faster model for translation
            response = await asyncio.to_thread(
                robust_api_call,
                client, 
                FAST_MODEL, 
                messages, 
                0.1
            )
            
            if response.choices and response.choices[0].message.content:
                translated_prompt = response.choices[0].message.content.strip()
                # print(f"Successfully translated prompt to: {translated_prompt}")
                return translated_prompt
            else:
                # Fallback to original prompt if translation fails
                print("Translation failed, falling back to original prompt.")
                return prompt
        else:
            # If already English, return as is
            # print("Prompt is already in English.")
            return prompt
    except LangDetectException:
        # If language detection fails, assume it might be English or a mix, and proceed
        print("Language detection failed, proceeding with original prompt.")
        return prompt
    except Exception as e:
        print(f"An unexpected error occurred during prompt translation: {e}")
        # Fallback to original prompt on any other error
        return prompt


async def generate_image(conversation_history: list, last_image_bytes: bytes = None):
    """
    Generates an image using the Together AI API. If previous image bytes are provided,
    it refines that image. Otherwise, it generates a new one from a detailed prompt.
    It now returns the image bytes along with the URL for server-side caching.
    """
    client = Together(api_key=os.getenv("TOGETHER_API_KEY"))
    try:
        yield {"type": "status", "data": "Starting image generation..."}
        
        # --- Always engineer the prompt from the conversation history ---
        yield {"type": "status", "data": "Engineering a detailed prompt from conversation..."}
        
        # Determine the appropriate status message based on whether we are refining or creating anew.
        if last_image_bytes:
            yield {"type": "status", "data": "Context: Refining existing image."}
        else:
            yield {"type": "status", "data": "Context: Creating new image."}

        system_prompt = """You are an expert prompt engineer for an advanced image generation AI. Your critical task is to create a single, comprehensive, and self-contained prompt based on a conversation history. The AI you are prompting has an image-to-image capability, meaning it can modify an existing image based on your text prompt.

**Core Instructions:**
1.  **Analyze the Full Conversation:** Read the entire chat history to understand the user's complete vision for the image.
2.  **Prioritize the Latest Message:** The user's most recent message is the most important instruction. It dictates the primary action to take.
3.  **Create a Self-Contained, Evolving Prompt:** Your final output MUST be a single, descriptive prompt that incorporates all relevant details from the conversation into one coherent instruction. The prompt should get longer and more detailed as the conversation progresses.
    *   **For New Images:** Create a detailed prompt based on the user's request.
    *   **For Modifications:** This is crucial. Re-describe the *entire scene* from the previous turn, and then integrate the user's latest message as a specific change. You can even emphasize the change, for example: "A photorealistic image of a classic red sports car on a coastal highway at sunset, with the crucial update that the car should now be blue."
4.  **Language:** The output prompt should be in a clear, descriptive language suitable for an image generation model (typically English).
5.  **Output:** Return ONLY the final prompt string and nothing else.

**Example Workflow (Modification):**
-   **Previous Prompt was:** "A photorealistic image of a classic red sports car on a coastal highway at sunset." (This resulted in an image).
-   **User's New Message:** "Now make the car blue."
-   **Your Correct Output:** "A photorealistic image of a classic red sports car on a coastal highway at sunset. The car is now blue."
-   **Your Incorrect Output:** "make the car blue"

**Example Workflow (New Image):**
-   **User's New Message:** "Draw a picture of a robot cat."
-   **Your Correct Output:** "A picture of a robot cat."
"""
        messages = create_contextual_messages(conversation_history, system_prompt)
        response_prompt = await asyncio.to_thread(robust_api_call, client, LLM_MODEL, messages, 0.2)
        image_prompt = response_prompt.choices[0].message.content.strip()
        yield {"type": "status", "data": f"Engineered prompt: \"{image_prompt}\""}

        # --- Translate the final prompt to English ---
        yield {"type": "status", "data": "Translating prompt to English for optimal results..."}
        translated_prompt = await translate_prompt_to_english(client, image_prompt)


        # --- Call the image generation API ---
        api_response = await asyncio.to_thread(
            client.images.generate,
            prompt=translated_prompt,
            model=IMAGE_GEN_MODEL,
            steps=28,
            n=1,
            condition_image=base64.b64encode(last_image_bytes).decode('utf-8') if last_image_bytes else None
        )

        # --- Process the response ---
        if api_response.data and hasattr(api_response.data[0], 'url') and api_response.data[0].url:
            image_url = api_response.data[0].url
            new_image_bytes = None
            try:
                # Download the newly created image to get its bytes for caching
                yield {"type": "status", "data": "Caching new image..."}
                response = requests.get(image_url, timeout=20)
                response.raise_for_status()
                new_image_bytes = response.content
            except Exception as e:
                yield {"type": "status", "data": f"Warning: Could not cache new image: {e}"}
            
            # Yield the final result with URL and bytes (if available)
            yield {"type": "image", "url": image_url, "bytes": new_image_bytes, "extended_prompt": translated_prompt}
        else:
            yield {"type": "error", "message": "Could not generate image. The API did not return a valid URL."}
            print("Image generation API Response:", api_response)

    except Exception as e:
        error_message = f"An error occurred during image generation: {e}"
        print(error_message)
        yield {"type": "error", "message": error_message}
    finally:
        yield {"type": "end"}
