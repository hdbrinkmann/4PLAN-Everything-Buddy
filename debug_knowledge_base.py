#!/usr/bin/env python3
"""
Debug-Script für Knowledge Base Probleme
Analysiert systematisch warum neue Dateien nicht in der Knowledge Base auftauchen
"""

import os
import json
import sys
from pathlib import Path
import docx
from datetime import datetime
import hashlib

# Pfad-Konfiguration
SCRIPT_DIR = os.path.dirname(os.path.abspath(__file__))
DOCUMENTS_PATH = os.path.join(SCRIPT_DIR, "Documents")
VECTOR_STORE_PATH = os.path.join(SCRIPT_DIR, "vector_store")
S4U_PATH = os.path.join(DOCUMENTS_PATH, "S4U & 4PLAN")

def print_header(title):
    """Formatierte Ausgabe von Überschriften"""
    print(f"\n{'='*60}")
    print(f"  {title}")
    print(f"{'='*60}")

def print_step(step_num, description):
    """Formatierte Ausgabe von Schritten"""
    print(f"\n[SCHRITT {step_num}] {description}")
    print("-" * 50)

def check_file_system():
    """Überprüft das Dateisystem und die Pfade"""
    print_step(1, "DATEISYSTEM-ANALYSE")
    
    # Basis-Pfade überprüfen
    print(f"Script-Verzeichnis: {SCRIPT_DIR}")
    print(f"Documents-Pfad: {DOCUMENTS_PATH}")
    print(f"Vector Store Pfad: {VECTOR_STORE_PATH}")
    print(f"S4U & 4PLAN Pfad: {S4U_PATH}")
    
    # Existenz prüfen
    paths_to_check = [
        ("Documents", DOCUMENTS_PATH),
        ("S4U & 4PLAN", S4U_PATH),
        ("Vector Store", VECTOR_STORE_PATH)
    ]
    
    for name, path in paths_to_check:
        if os.path.exists(path):
            print(f"✓ {name}: Existiert")
            if os.path.isdir(path):
                print(f"  - Ist Verzeichnis: Ja")
                print(f"  - Berechtigung lesbar: {os.access(path, os.R_OK)}")
                print(f"  - Berechtigung schreibbar: {os.access(path, os.W_OK)}")
            else:
                print(f"  - Ist Verzeichnis: Nein (Problem!)")
        else:
            print(f"✗ {name}: Existiert nicht (Problem!)")

def analyze_files_in_s4u():
    """Analysiert alle Dateien im S4U & 4PLAN Ordner"""
    print_step(2, "DATEI-ANALYSE IM S4U & 4PLAN ORDNER")
    
    if not os.path.exists(S4U_PATH):
        print("✗ S4U & 4PLAN Ordner existiert nicht!")
        return []
    
    # Alle Dateien auflisten
    all_files = []
    for root, dirs, files in os.walk(S4U_PATH):
        for file in files:
            full_path = os.path.join(root, file)
            all_files.append(full_path)
    
    print(f"Gefundene Dateien gesamt: {len(all_files)}")
    
    # Nur .docx und .pdf Dateien filtern
    document_files = []
    for file_path in all_files:
        if file_path.lower().endswith(('.docx', '.pdf')):
            # Temporäre Dateien ausschließen
            if not os.path.basename(file_path).startswith('~$'):
                document_files.append(file_path)
    
    print(f"Relevante Dokumente (.docx/.pdf): {len(document_files)}")
    
    # Detailierte Analyse jeder Datei
    for i, file_path in enumerate(document_files, 1):
        print(f"\n[DATEI {i}] {os.path.basename(file_path)}")
        
        # Dateigröße
        file_size = os.path.getsize(file_path)
        print(f"  - Größe: {file_size / (1024*1024):.2f} MB")
        
        # Letztes Änderungsdatum
        mtime = os.path.getmtime(file_path)
        mod_time = datetime.fromtimestamp(mtime).strftime("%Y-%m-%d %H:%M:%S")
        print(f"  - Letzte Änderung: {mod_time}")
        
        # Datei-Hash für Eindeutigkeit
        with open(file_path, 'rb') as f:
            file_hash = hashlib.md5(f.read()).hexdigest()[:8]
        print(f"  - Hash: {file_hash}")
        
        # Versuchen, .docx Datei zu öffnen
        if file_path.lower().endswith('.docx'):
            try:
                doc = docx.Document(file_path)
                paragraph_count = len(doc.paragraphs)
                text_paragraphs = [p for p in doc.paragraphs if p.text.strip()]
                print(f"  - Absätze gesamt: {paragraph_count}")
                print(f"  - Absätze mit Text: {len(text_paragraphs)}")
                
                # Erste paar Zeilen Text extrahieren
                first_text = ""
                for p in text_paragraphs[:3]:
                    if p.text.strip():
                        first_text += p.text.strip() + " "
                
                if first_text:
                    preview = first_text[:100] + "..." if len(first_text) > 100 else first_text
                    print(f"  - Textvorschau: {preview}")
                else:
                    print("  - ✗ Kein Text gefunden (Problem!)")
                    
            except Exception as e:
                print(f"  - ✗ Fehler beim Öffnen: {e}")
        
        print(f"  - Voller Pfad: {file_path}")
    
    return document_files

def test_get_document_list():
    """Testet die get_document_list Funktion aus llm.py"""
    print_step(3, "TEST DER get_document_list() FUNKTION")
    
    try:
        # Importiere die Funktion
        sys.path.insert(0, SCRIPT_DIR)
        from llm import get_document_list
        
        # Teste die Funktion
        found_files = get_document_list(S4U_PATH)
        
        print(f"Von get_document_list() gefundene Dateien: {len(found_files)}")
        
        for i, file_path in enumerate(found_files, 1):
            print(f"  [{i}] {os.path.basename(file_path)}")
            print(f"      Pfad: {file_path}")
        
        if not found_files:
            print("✗ get_document_list() hat keine Dateien gefunden! (Problem!)")
        
        return found_files
        
    except Exception as e:
        print(f"✗ Fehler beim Test der get_document_list() Funktion: {e}")
        return []

def test_docx_processing():
    """Testet die DOCX-Verarbeitung für eine spezifische Datei"""
    print_step(4, "TEST DER DOCX-VERARBEITUNG")
    
    try:
        sys.path.insert(0, SCRIPT_DIR)
        from llm import smart_chunk_document
        
        # Finde die neueste .docx Datei
        docx_files = [f for f in os.listdir(S4U_PATH) if f.lower().endswith('.docx') and not f.startswith('~$')]
        
        if not docx_files:
            print("✗ Keine .docx Dateien gefunden zum Testen!")
            return
        
        # Sortiere nach Änderungsdatum (neueste zuerst)
        docx_files.sort(key=lambda f: os.path.getmtime(os.path.join(S4U_PATH, f)), reverse=True)
        
        test_file = os.path.join(S4U_PATH, docx_files[0])
        print(f"Teste Datei: {docx_files[0]}")
        
        # Verarbeite die Datei
        chunks = smart_chunk_document(test_file)
        
        print(f"Generierte Chunks: {len(chunks)}")
        
        # Analysiere die ersten paar Chunks
        for i, chunk in enumerate(chunks[:3], 1):
            print(f"\n[CHUNK {i}]")
            print(f"  - Länge: {len(chunk.page_content)} Zeichen")
            print(f"  - Quelle: {chunk.metadata.get('source', 'N/A')}")
            print(f"  - Überschrift: {chunk.metadata.get('heading', 'N/A')}")
            print(f"  - Knowledge Field: {chunk.metadata.get('knowledge_field', 'N/A')}")
            
            # Vorschau des Inhalts
            preview = chunk.page_content[:200] + "..." if len(chunk.page_content) > 200 else chunk.page_content
            print(f"  - Inhalt: {preview}")
        
        if len(chunks) == 0:
            print("✗ Keine Chunks generiert! (Problem!)")
        
        return chunks
        
    except Exception as e:
        print(f"✗ Fehler beim Test der DOCX-Verarbeitung: {e}")
        import traceback
        traceback.print_exc()
        return []

def check_vector_store():
    """Überprüft den Vector Store Status"""
    print_step(5, "VECTOR STORE ANALYSE")
    
    if not os.path.exists(VECTOR_STORE_PATH):
        print("✗ Vector Store Pfad existiert nicht!")
        return
    
    # Unterordner auflisten
    subdirs = [d for d in os.listdir(VECTOR_STORE_PATH) if os.path.isdir(os.path.join(VECTOR_STORE_PATH, d))]
    
    print(f"Vector Store Unterordner: {len(subdirs)}")
    
    for subdir in subdirs:
        subdir_path = os.path.join(VECTOR_STORE_PATH, subdir)
        print(f"\n[VECTOR STORE] {subdir}")
        
        # Dateien im Unterordner
        files = os.listdir(subdir_path)
        print(f"  - Dateien: {files}")
        
        # Prüfe auf FAISS Index
        faiss_index = os.path.join(subdir_path, "index.faiss")
        pkl_file = os.path.join(subdir_path, "index.pkl")
        
        if os.path.exists(faiss_index):
            index_size = os.path.getsize(faiss_index)
            print(f"  - FAISS Index: Existiert ({index_size} Bytes)")
        else:
            print(f"  - FAISS Index: ✗ Nicht vorhanden")
        
        if os.path.exists(pkl_file):
            pkl_size = os.path.getsize(pkl_file)
            print(f"  - PKL Datei: Existiert ({pkl_size} Bytes)")
        else:
            print(f"  - PKL Datei: ✗ Nicht vorhanden")

def check_knowledge_fields():
    """Überprüft die knowledge_fields.json Datei"""
    print_step(6, "KNOWLEDGE FIELDS KONFIGURATION")
    
    fields_file = os.path.join(SCRIPT_DIR, "knowledge_fields.json")
    
    if os.path.exists(fields_file):
        print(f"✓ knowledge_fields.json existiert")
        
        try:
            with open(fields_file, 'r') as f:
                fields = json.load(f)
            
            print(f"Konfigurierte Knowledge Fields: {len(fields)}")
            for i, field in enumerate(fields, 1):
                print(f"  [{i}] {field}")
            
            # Prüfe ob "S4U & 4PLAN" dabei ist
            if "S4U & 4PLAN" in fields:
                print("✓ 'S4U & 4PLAN' ist in den Knowledge Fields konfiguriert")
            else:
                print("✗ 'S4U & 4PLAN' fehlt in den Knowledge Fields! (Problem!)")
                
        except Exception as e:
            print(f"✗ Fehler beim Lesen der knowledge_fields.json: {e}")
    else:
        print("✗ knowledge_fields.json existiert nicht!")

def generate_test_summary():
    """Generiert eine Zusammenfassung der Tests"""
    print_header("DIAGNOSE-ZUSAMMENFASSUNG")
    
    print("Durchgeführte Tests:")
    print("1. ✓ Dateisystem-Analyse")
    print("2. ✓ Datei-Analyse im S4U & 4PLAN Ordner")
    print("3. ✓ Test der get_document_list() Funktion")
    print("4. ✓ Test der DOCX-Verarbeitung")
    print("5. ✓ Vector Store Analyse")
    print("6. ✓ Knowledge Fields Konfiguration")
    
    print("\nNächste Schritte:")
    print("- Probleme basierend auf den Ergebnissen identifizieren")
    print("- Lösungsvorschläge implementieren")
    print("- Knowledge Base Update erneut testen")

def main():
    """Haupt-Funktion"""
    print_header("KNOWLEDGE BASE DEBUG-ANALYSE")
    print(f"Analysezeitpunkt: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
    
    try:
        # Tests ausführen
        check_file_system()
        found_files = analyze_files_in_s4u()
        get_document_list_result = test_get_document_list()
        chunks = test_docx_processing()
        check_vector_store()
        check_knowledge_fields()
        
        generate_test_summary()
        
    except Exception as e:
        print(f"\n✗ Kritischer Fehler während der Analyse: {e}")
        import traceback
        traceback.print_exc()

if __name__ == "__main__":
    main()
